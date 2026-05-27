"""
AUSMAR PSE QA Engine — Production Review Pipeline (v2: Content-Based Classification)

Documents are identified by their CONTENT (PDF text, spreadsheet structure, image analysis),
NOT by filename. This eliminates false positives from inconsistent naming by consultants.

Flow:
1. Extract files from zip, remove junk
2. Extract text/content from each file (pypdf for PDFs, openpyxl for spreadsheets)
3. Classify each file by content using keyword matching, then LLM fallback for ambiguous files
4. Build a file_map (doc_type -> file) used by all downstream checks
5. Rename files to standard names based on classification
6. Run completeness check against the classified map
7. Run vision checks (GeoSite, Red Pen) using the classified file map
8. Generate verdict
"""

import os
import re
import json
import zipfile
import shutil
import base64
import tempfile
import traceback
import gc
from datetime import datetime
from pathlib import Path

from openai import OpenAI
from PIL import Image
import io

import database as db

# PDF text extraction
try:
    from pypdf import PdfReader
    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False

# Spreadsheet inspection
try:
    import openpyxl
    _OPENPYXL_AVAILABLE = True
except ImportError:
    _OPENPYXL_AVAILABLE = False

# PDF-to-image for vision analysis
_PDF2IMAGE_AVAILABLE = True
try:
    from pdf2image import convert_from_path
except ImportError:
    _PDF2IMAGE_AVAILABLE = False


# ---------------------------------------------------------------------------
# Pre-Log Parser
# ---------------------------------------------------------------------------
def parse_prelog(notes: str) -> dict:
    """Parse Pre-Log notes to extract key fields.
    
    Looks for patterns like:
    - Land Registered (Y/N): NO
    - NHP or STC: STC
    - Estate & Stage: Aura P18 Stage 3
    """
    prelog = {
        "land_registered": None,  # True/False/None
        "nhp_or_stc": None,  # "NHP" or "STC"
        "estate_stage": None,
    }
    
    if not notes:
        return prelog
    
    # Land Registered (Y/N)
    match = re.search(r'Land Registered\s*\(Y/N\)\s*:\s*([YN])', notes, re.IGNORECASE)
    if match:
        prelog["land_registered"] = match.group(1).upper() == 'Y'
    
    # NHP or STC
    match = re.search(r'NHP or STC\s*:\s*(NHP|STC)', notes, re.IGNORECASE)
    if match:
        prelog["nhp_or_stc"] = match.group(1).upper()
    
    # Estate & Stage
    match = re.search(r'Estate & Stage\s*:\s*([^\n]+)', notes)
    if match:
        prelog["estate_stage"] = match.group(1).strip()
    
    return prelog


# ---------------------------------------------------------------------------
# OpenAI client (lazy init)
# ---------------------------------------------------------------------------
_client = None

def get_client() -> OpenAI:
    global _client
    if _client is None:
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY environment variable is not set.")
        base_url = os.environ.get("OPENAI_BASE_URL")
        if base_url:
            _client = OpenAI(api_key=api_key, base_url=base_url)
        else:
            _client = OpenAI(api_key=api_key)
    return _client


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
JUNK_PATTERNS = {"__macosx", ".ds_store", "thumbs.db", ".tmp", "~$"}
LICENCE_BAD_EXTS = {".heic", ".msg", ".webp"}
MAX_IMAGE_DIM = 1024

# NOTE: Gas cooktops are permitted in all estates including those with gas-ban covenants.
# Stockland ruling overrides covenant restrictions. Do NOT flag gas cooktops.

# Standard document types and their canonical filenames
DOC_TYPES = {
    "pse_doc": "PSE Doc (Signed)",
    "pse_excel": "PSE Excel",
    "geosite": "GeoSite Plan (Signed)",
    "itp": "ITP Form (Signed)",
    "deposit_receipt": "Deposit Receipt",
    "drivers_licence": "Drivers Licence",
    "pse_checklist": "PSE Checklist",
    "red_pen": "Red Pen Markup (Signed)",
    "pod_envelope": "POD or Building Envelope",
    "compaction_report": "Compaction Report",
    "covenant_guidelines": "Covenant Design Guidelines",
    "disclosure_plan": "Disclosure Plan",
    "covenant_application": "Covenant Application (Signed)",
    "pool_form": "Pool Form (Signed)",
    "discount_approval": "Discount Approval",
    "owner_supplied": "Owner Supplied Items Approval",
    "modified_plan": "Modified Plan Approval",
    "promo_ack": "Promo Client Acknowledgement (Signed)",
    "fall_ack": "Sites with Fall Acknowledgment (Signed)",
    "acoustic_report": "Acoustic Report",
    "bal_report": "BAL Report",
    "contour_survey": "Contour Survey",
    "sales_accept": "Sales Accept Doc",
    "soil_report": "Soil Report",
    "covenant_doc": "Covenant",
}

# Core required documents (missing = issue)
# Per 1.0 PSE Document Naming: PSE Signed, PSE Excel, GeoSite, ITP, Deposit Receipt, Drivers Licence,
# Red Pen Markup, Promo Ack are all required for a PSE submission.
CORE_REQUIRED = ["pse_doc", "pse_excel", "geosite", "itp", "deposit_receipt", "drivers_licence", "red_pen", "promo_ack"]
# Soft required (missing = warning)
SOFT_REQUIRED = ["pse_checklist"]
# Conditionally expected (missing = warning if not explained)
# NOTE: Disclosure Plan, POD, Building Envelope are NOT flagged as missing — they are
# job-specific (only required for registered lots) and the team knows when to include them.
CONDITIONAL_EXPECTED = []


# ---------------------------------------------------------------------------
# Regional Compliance Rules
# ---------------------------------------------------------------------------
# Per AUSMAR compliance checklists for each council/estate area.
# Detection keywords are matched against estate name, address, or council area
# extracted from the GeoSite plan or ITP form.
REGIONAL_RULES = {
    "SCRC": {
        "name": "Sunshine Coast Regional Council",
        "detection_keywords": [
            "pelican waters", "bokarina", "birtinya", "sippy downs", "mountain creek",
            "buderim", "caloundra", "kawana", "mooloolaba", "sunshine coast regional",
            "scrc", "maroochydore", "mudjimba", "bli bli", "coolum", "peregian",
        ],
        "lhdc_required": False,
        "setbacks_to": "WALL",
        "max_site_coverage_pct": 60,
        "front_setback_m": 3.0,
        "garage_setback_m": 5.5,
        "side_setback_m": 1.0,
        "rear_setback_m": 1.5,
        "max_build_height_m": 8.5,
        "crossover_width_max_m": 4.8,
        "site_fall_threshold_mm": 300,
    },
    "MBRC": {
        "name": "Moreton Bay Regional Council",
        "detection_keywords": [
            "ridgeview", "moreton bay", "caboolture", "narangba", "burpengary",
            "morayfield", "petrie", "strathpine", "kallangur", "mbrc",
            "north lakes", "mango hill", "griffin", "dakabin",
        ],
        "lhdc_required": False,
        "setbacks_to": "WALL",
        "max_site_coverage_pct": 75,
        "front_setback_m": 3.0,
        "garage_setback_m": 5.4,
        "side_setback_m": 0.75,
        "rear_setback_m": 0.75,
        "max_build_height_m": 8.5,
        "roof_pitch_max_deg": 25,
        "eave_depth_mm": 450,
        "porch_min_depth_m": 1.5,
        "max_facade_material_pct": 80,
        "garage_setback_from_fbl_mm": 500,
        "crossover_width_max_m": 4.8,
        "landscaping_strip_min_mm": 500,
        "side_fencing_return_m": 1.0,
        "site_fall_threshold_mm": 500,
        "fall_over_build_pad_threshold_mm": 300,
    },
    "NOOSA": {
        "name": "Noosa Council",
        "detection_keywords": [
            "noosa", "tewantin", "cooroy", "pomona", "noosaville",
            "sunrise beach", "peregian beach", "marcus beach",
        ],
        "lhdc_required": True,  # Always required regardless of frontage
        "setbacks_to": "OMP",
        "max_site_coverage_pct": 40,
        "front_setback_m": 6.0,
        "side_setback_m": 1.5,  # Up to 4.5m high
        "rear_setback_m": 3.0,  # Up to 4.5m high
        "max_build_height_m": 8.0,
        "crossover_width_max_m": 6.0,
        "site_fall_threshold_mm": 3000,
        "fall_over_build_pad_threshold_mm": 2000,
    },
    "AURA_PV300": {
        "name": "Aura (PV300 - Rivermont, larger lots)",
        "detection_keywords": [
            "rivermont", "aura", "baringa", "nirimba", "harmony",
            "potential impact buffer", "emerging community",
        ],
        "lhdc_required": True,
        "setbacks_to": "WALL",
        "max_site_coverage_pct": 60,
        "front_setback_m": 3.0,
        "garage_setback_m": 5.4,
        "side_setback_m": 1.05,
        "rear_setback_m": 1.05,
        "max_build_height_m": 9.0,
        "roof_pitch_min_deg": 22.5,
        "eave_return_min_m": 2.0,
        "eave_depth_mm": 450,
        "porch_min_depth_m": 2.0,
        "porch_min_area_m2": 4.0,
        "max_facade_material_pct": 60,
        "garage_setback_from_fbl_mm": 450,
        "ev_charging_required": True,
        "crossover_width_max_m": 4.8,
        "landscaping_strip_min_mm": 500,
        "side_rear_fencing_height_m": 1.8,
        "pos_min_area_m2": 20,
        "pos_min_dimension_m": 4.0,
        "site_fall_threshold_mm": 800,
        "fall_over_build_pad_threshold_mm": 300,
    },
    "AURA_V300": {
        "name": "Aura (V300 - Acacia, small lots)",
        "detection_keywords": [
            "acacia", "aura", "baringa", "nirimba",
        ],
        "lhdc_required": False,
        "setbacks_to": "WALL",
        "max_site_coverage_pct": 60,
        "front_setback_m": 2.4,
        "garage_setback_m": 5.0,
        "side_setback_m": 0.9,
        "rear_setback_m": 1.5,
        "max_build_height_m": 9.0,
        "roof_pitch_max_deg": 30,
        "eave_return_min_m": 2.0,
        "eave_depth_mm": 450,
        "porch_min_depth_m": 1.5,
        "porch_min_area_m2": 3.0,
        "max_facade_material_pct": 80,
        "garage_setback_from_fbl_m": 1.0,
        "ev_charging_required": True,
        "crossover_width_max_m": 3.0,
        "landscaping_strip_min_mm": 500,
        "landscaping_strip_max_m": 1.0,
        "water_tank_litres": 3000,
        "side_fencing_return_m": 1.0,
        "side_rear_fencing_height_m": 1.8,
        "pos_min_area_m2": 15,
        "pos_min_dimension_m": 3.0,
        "site_fall_threshold_mm": 480,
        "fall_over_build_pad_threshold_mm": 250,
    },
    "GYMPIE": {
        "name": "Gympie Regional Council",
        "detection_keywords": [
            "gympie", "bellagrove", "rural estate", "bal zone", "gympie regional",
        ],
        "lhdc_required": True,
        "setbacks_to": "OMP",
        "max_site_coverage_pct": 50,
        "front_setback_m": 6.0,
        "side_setback_m": 1.5,
        "uf_side_setback_m": 2.0,
        "rear_setback_gf_m": 1.5,
        "rear_setback_ff_m": 2.0,
        "max_build_height_m": 8.5,
        "crossover_width_max_m": 6.0,
        "landscaping_strip_min_m": 1.0,
        "site_fall_threshold_mm": 2500,
        "fall_over_build_pad_threshold_mm": 1000,
    },
    "KINMA": {
        "name": "Kinma Valley",
        "detection_keywords": [
            "kinma valley", "kinma",
        ],
        "lhdc_required": False,
        "setbacks_to": "BOTH",
        "max_site_coverage_pct": 60,
        "front_setback_omp_m": 2.0,
        "front_setback_wall_m": 3.0,
        "garage_setback_m": 5.4,
        "side_setback_wall_m": 1.0,
        "side_setback_omp_m": 0.45,
        "rear_setback_wall_m": 1.0,
        "rear_setback_omp_m": 0.45,
        "max_build_height_m": 8.5,
        "roof_pitch_min_deg": 20,
        "eave_return_min_m": 1.5,
        "eave_depth_min_mm": 450,
        "max_facade_material_pct": 80,
        "garage_setback_from_fbl_mm": 500,
        "crossover_width_max_m": 5.0,
        "side_fencing_return_m": 1.0,
        "side_rear_fencing_height_m": 1.8,
        "site_fall_threshold_mm": 1000,
        "fall_over_build_pad_threshold_mm": 500,
    },
}


def detect_regional_rules(estate_name: str, address: str, council: str = "") -> dict | None:
    """Detect which regional compliance rules apply based on estate/address/council keywords.
    Returns the matching rules dict or None if no match."""
    search_text = f"{estate_name} {address} {council}".lower()
    
    # Check AURA variants first (more specific) — V300 is for small lots (Acacia)
    if "acacia" in search_text and "aura" in search_text:
        return REGIONAL_RULES["AURA_V300"]
    if "rivermont" in search_text or ("aura" in search_text and "potential impact" in search_text):
        return REGIONAL_RULES["AURA_PV300"]
    
    # Check each region
    for region_key, rules in REGIONAL_RULES.items():
        if region_key.startswith("AURA"):
            continue  # Already handled above
        for keyword in rules["detection_keywords"]:
            if keyword in search_text:
                return rules
    
    # Check AURA generically (if no specific variant matched)
    if "aura" in search_text:
        return REGIONAL_RULES["AURA_PV300"]  # Default to PV300 if just "aura"
    
    return None


def should_apply_lhdc(lot_frontage_m: float | None, regional_rules: dict | None) -> tuple[bool, str]:
    """Determine if LHDC checks should apply.
    Returns (apply_lhdc: bool, reason: str).
    
    Rules:
    - Frontage <= 12.5m: EXEMPT from LHDC regardless of council
    - NOOSA, GYMPIE, AURA PV300: Always require LHDC (even if frontage > 12.5m)
    - Other councils: LHDC not required
    """
    # Frontage exemption — 12.5m or less is always exempt
    if lot_frontage_m is not None and lot_frontage_m <= 12.5:
        return False, f"Lot frontage {lot_frontage_m}m is ≤12.5m — LHDC exempt"
    
    if regional_rules is None:
        # No regional rules detected — apply LHDC conservatively
        if lot_frontage_m is not None and lot_frontage_m > 12.5:
            return True, f"Lot frontage {lot_frontage_m}m is >12.5m — LHDC may apply"
        # Frontage unknown — apply LHDC conservatively (better to check than miss)
        return True, "Lot frontage unknown — applying LHDC check conservatively (verify frontage)"
    
    if regional_rules.get("lhdc_required") is True:
        return True, f"LHDC required per {regional_rules['name']} requirements"
    
    return False, f"LHDC not required per {regional_rules['name']} requirements"


# ---------------------------------------------------------------------------
# Text extraction utilities
# ---------------------------------------------------------------------------
def extract_pdf_text(pdf_path: str, max_pages: int = 5) -> str:
    """Extract text from a PDF using pypdf. Returns empty string on failure."""
    if not _PYPDF_AVAILABLE:
        return ""
    try:
        reader = PdfReader(pdf_path)
        text_parts = []
        for i, page in enumerate(reader.pages[:max_pages]):
            t = page.extract_text() or ""
            text_parts.append(t)
        return "\n".join(text_parts)[:8000]  # Cap at 8K chars to save memory
    except Exception as e:
        print(f"[WARN] extract_pdf_text failed for {pdf_path}: {e}")
        return ""


def extract_spreadsheet_info(file_path: str) -> dict:
    """Extract sheet names and header row from xlsx/xlsm/xlsb. Returns dict with metadata."""
    ext = Path(file_path).suffix.lower()
    # Handle .xlsb (binary Excel) — openpyxl cannot read it, use pyxlsb if available
    if ext == ".xlsb":
        try:
            import pyxlsb
            with pyxlsb.open_workbook(file_path) as wb:
                sheets = wb.sheets
                sample_values = []
                for sheet_name in sheets[:3]:
                    with wb.get_sheet(sheet_name) as ws:
                        row_count = 0
                        for row in ws.rows():
                            row_count += 1
                            if row_count > 10:
                                break
                            row_text = [str(c.v).lower() for c in row if c.v is not None]
                            sample_values.extend(row_text)
                return {
                    "type": "spreadsheet",
                    "sheets": sheets,
                    "headers": sample_values[:20],
                    "sample_text": " ".join(sample_values)[:4000],
                }
        except ImportError:
            # pyxlsb not available — classify by filename alone
            fname_lower = Path(file_path).name.lower()
            sample = fname_lower  # Use filename as sample text for PSE detection
            return {"type": "spreadsheet", "sheets": [], "headers": [], "sample_text": sample}
        except Exception as e:
            print(f"[WARN] extract_spreadsheet_info (xlsb) failed for {file_path}: {e}")
            fname_lower = Path(file_path).name.lower()
            return {"type": "spreadsheet", "sheets": [], "headers": [], "sample_text": fname_lower}
    if not _OPENPYXL_AVAILABLE:
        return {"type": "spreadsheet", "sheets": [], "headers": [], "sample_text": ""}
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets = wb.sheetnames
        headers = []
        sample_values = []
        for sheet_name in sheets[:3]:
            ws = wb[sheet_name]
            row_count = 0
            for row in ws.iter_rows(max_row=10, values_only=True):
                row_count += 1
                row_text = [str(c).lower() for c in row if c is not None]
                if row_count <= 2:
                    headers.extend(row_text)
                sample_values.extend(row_text)
        wb.close()
        return {
            "type": "spreadsheet",
            "sheets": sheets,
            "headers": headers,
            "sample_text": " ".join(sample_values)[:4000],
        }
    except Exception as e:
        print(f"[WARN] extract_spreadsheet_info failed for {file_path}: {e}")
        return {"type": "spreadsheet", "sheets": [], "headers": [], "sample_text": ""}


def extract_pse_excel_totals(file_path: str, original_name: str = "") -> dict:
    """Extract TOTAL DEBITS, TOTAL CREDITS, and PROPOSED NEW HOME PRICE from PSE Excel.
    Also extracts the pricing period from the filename or sheet data.
    Returns dict with keys: total_debits, total_credits, proposed_price, pricing_period."""
    result = {"total_debits": None, "total_credits": None, "proposed_price": None, "pricing_period": None}
    if not _OPENPYXL_AVAILABLE:
        return result
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        # Search the PSE-NHP sheet (main pricing sheet) or GLD PSE
        target_sheets = [s for s in wb.sheetnames if s.lower() in ["pse-nhp", "gld pse", "pse"]]
        if not target_sheets:
            # Fall back to any sheet with "pse" in the name
            target_sheets = [s for s in wb.sheetnames if "pse" in s.lower()]
        if not target_sheets:
            target_sheets = wb.sheetnames[:1]  # Last resort: first sheet

        for sheet_name in target_sheets:
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                if row is None:
                    continue
                # Look for label in column index 1 (second column)
                row_vals = list(row)
                label = None
                for i, cell in enumerate(row_vals):
                    if isinstance(cell, str):
                        label_lower = cell.strip().lower()
                        if "total debits" in label_lower:
                            # Value is typically in column index 8 or the next numeric column
                            for v in row_vals[i+1:]:
                                if isinstance(v, (int, float)) and v > 0:
                                    result["total_debits"] = round(v, 2)
                                    break
                        elif "total credits" in label_lower:
                            for v in row_vals[i+1:]:
                                if isinstance(v, (int, float)) and v > 0:
                                    result["total_credits"] = round(v, 2)
                                    break
                        elif "proposed new home price" in label_lower:
                            for v in row_vals[i+1:]:
                                if isinstance(v, (int, float)) and v > 0:
                                    result["proposed_price"] = round(v, 2)
                                    break
            # If we found the proposed price, stop searching other sheets
            if result["proposed_price"] is not None:
                break

        wb.close()
    except Exception as e:
        print(f"[WARN] extract_pse_excel_totals failed: {e}")

    # Try to extract pricing period from filename
    # Use original_name if available (before renaming), otherwise use the file path basename
    fname = (original_name or os.path.basename(file_path)).lower()
    # Common patterns: "PSE Dec-Jan.xlsm", "PSE Feb-Mar 2026.xlsx"
    month_names = ["jan", "feb", "mar", "apr", "may", "jun",
                   "jul", "aug", "sep", "oct", "nov", "dec"]
    # Find months by their POSITION in the filename (not iteration order)
    found_months = []
    for m in month_names:
        pos = fname.find(m)
        if pos >= 0:
            found_months.append((pos, m))
    # Sort by position in filename to preserve the actual order (e.g., Dec-Jan not Jan-Dec)
    found_months.sort(key=lambda x: x[0])
    month_list = [m for _, m in found_months]
    if len(month_list) >= 2:
        result["pricing_period"] = f"{month_list[0].title()}-{month_list[1].title()}"
    elif len(month_list) == 1:
        result["pricing_period"] = month_list[0].title()

    return result


# ---------------------------------------------------------------------------
# Image/PDF conversion utilities (for vision analysis)
# ---------------------------------------------------------------------------
def _resize_image(img: Image.Image, max_dim: int = MAX_IMAGE_DIM) -> Image.Image:
    w, h = img.size
    if w <= max_dim and h <= max_dim:
        return img
    if w > h:
        new_w = max_dim
        new_h = int(h * max_dim / w)
    else:
        new_h = max_dim
        new_w = int(w * max_dim / h)
    return img.resize((new_w, new_h), Image.LANCZOS)


def _img_to_b64(img: Image.Image) -> str:
    img = _resize_image(img)
    if img.mode != "RGB":
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=75)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    buf.close()
    return b64


def pdf_page_to_base64(pdf_path: str, page_num: int = 0, dpi: int = 100) -> str:
    if not _PDF2IMAGE_AVAILABLE:
        return ""
    try:
        images = convert_from_path(pdf_path, first_page=page_num + 1, last_page=page_num + 1, dpi=dpi)
        if not images:
            return ""
        b64 = _img_to_b64(images[0])
        for img in images:
            img.close()
        del images
        gc.collect()
        return b64
    except Exception as e:
        print(f"[WARN] pdf_page_to_base64 failed for {pdf_path} page {page_num}: {e}")
        return ""


def pdf_all_pages_to_base64(pdf_path: str, dpi: int = 100, max_pages: int = 3) -> list[str]:
    if not _PDF2IMAGE_AVAILABLE:
        return []
    try:
        images = convert_from_path(pdf_path, dpi=dpi, last_page=max_pages)
        results = []
        for img in images[:max_pages]:
            results.append(_img_to_b64(img))
            img.close()
        del images
        gc.collect()
        return results
    except Exception as e:
        print(f"[WARN] pdf_all_pages_to_base64 failed for {pdf_path}: {e}")
        return []


def image_to_base64(img_path: str) -> str:
    try:
        with Image.open(img_path) as img:
            return _img_to_b64(img)
    except Exception as e:
        print(f"[WARN] image_to_base64 failed for {img_path}: {e}")
        return ""


# ---------------------------------------------------------------------------
# LLM helpers
# ---------------------------------------------------------------------------
def call_text_model(system_prompt: str, user_prompt: str, model: str = "gpt-4.1-nano") -> str:
    resp = get_client().chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
        max_tokens=4000,
    )
    return resp.choices[0].message.content.strip()


def call_vision_model(system_prompt: str, user_text: str, image_b64_list: list[str], model: str = "gpt-4.1-mini") -> str:
    content = [{"type": "text", "text": user_text}]
    for b64 in image_b64_list:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "low"},
        })
    resp = get_client().chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content},
        ],
        temperature=0.1,
        max_tokens=4000,
    )
    return resp.choices[0].message.content.strip()


def parse_json_from_llm(raw: str) -> dict:
    json_match = re.search(r"\{[\s\S]*\}", raw)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError:
            pass
    return {"raw_response": raw}


# ---------------------------------------------------------------------------
# CONTENT-BASED DOCUMENT CLASSIFIER
# ---------------------------------------------------------------------------
# Text fingerprints: keywords/phrases that identify each document type from content.
# Checked against extracted PDF text or spreadsheet content.
# Order matters — more specific patterns first to avoid misclassification.
CONTENT_FINGERPRINTS = [
    # PSE Doc — the main Provisional Sales Estimate PDF
    ("pse_doc", {
        "keywords": ["provisional sales estimate", "total contract price", "base price",
                      "inclusions and exclusions", "pse document", "sales estimate"],
        "min_matches": 1,
        "extensions": [".pdf", ".docx"],
    }),
    # ITP — Intention to Purchase agreement
    ("itp", {
        "keywords": ["intention to purchase", "purchaser details", "purchaser name",
                      "intention to proceed", "purchase agreement", "itp form"],
        "min_matches": 1,
        "extensions": [".pdf", ".docx"],
    }),
    # GeoSite Plan
    ("geosite", {
        "keywords": ["geosite", "geo site", "geosite.com.au", "site plan",
                      "setback", "site coverage", "building envelope"],
        "min_matches": 2,  # Need 2+ because "setback" alone is too generic
        "extensions": [".pdf", ".jpg", ".jpeg", ".png"],
    }),
    # Red Pen Markup
    ("red_pen", {
        "keywords": ["red pen", "markup", "floor plan", "elevations", "electrical plan",
                      "floor coverings", "concrete plan"],
        "min_matches": 2,
        "extensions": [".pdf", ".jpg", ".jpeg", ".png"],
    }),
    # PSE Checklist — requires specific PSE checklist language, not generic checklist words
    ("pse_checklist", {
        "keywords": ["pse checklist", "document checklist", "1.0 pse document naming",
                      "pse documents list", "document list", "document naming"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Deposit Receipt
    ("deposit_receipt", {
        "keywords": ["deposit receipt", "receipt", "payment received", "deposit amount",
                      "$2,500", "$4,000", "$2500", "$4000", "preliminary deposit"],
        "min_matches": 1,
        "extensions": [".pdf", ".jpg", ".jpeg", ".png"],
    }),
    # Covenant Application
    ("covenant_application", {
        "keywords": ["covenant application", "design approval application",
                      "design review application", "covenant form"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Pool Form
    ("pool_form", {
        "keywords": ["swimming pool", "pool form", "pool fencing",
                      "pool barrier", "pool safety"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Sites with Fall Acknowledgment
    ("fall_ack", {
        "keywords": ["sites with fall", "fall acknowledgment", "fall acknowledgement",
                      "cut and fill", "retaining wall acknowledgment"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Promo Client Acknowledgement
    ("promo_ack", {
        "keywords": ["client acknowledgement", "promotion acknowledgement",
                      "advantage", "super saver", "promo acknowledgement",
                      "promotional offer"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Acoustic Report — min_matches=2 to avoid beating Disclosure Plan on single keyword
    ("acoustic_report", {
        "keywords": ["acoustic", "noise assessment", "acoustic report",
                      "noise category", "acoustic category", "db(a)"],
        "min_matches": 2,
        "extensions": [".pdf"],
    }),
    # BAL Report
    ("bal_report", {
        "keywords": ["bushfire attack level", "bal report", "bal assessment",
                      "bushfire prone", "bal-"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Compaction Report
    ("compaction_report", {
        "keywords": ["compaction", "bearing capacity", "compaction test",
                      "soil compaction", "density ratio"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Soil Report
    ("soil_report", {
        "keywords": ["soil test", "geotechnical", "soil classification",
                      "site classification", "foundation recommendation",
                      "reactive soil"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Contour Survey
    ("contour_survey", {
        "keywords": ["contour survey", "topographic survey", "contour plan",
                      "reduced levels", "spot levels", "topographical"],
        "min_matches": 1,
        "extensions": [".pdf", ".jpg", ".jpeg", ".png"],
    }),
    # Disclosure Plan / Survey Plan
    ("disclosure_plan", {
        "keywords": ["disclosure plan", "plan of subdivision", "survey plan",
                      "plan of survey", "disclosure statement", "community title"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # POD or Building Envelope
    ("pod_envelope", {
        "keywords": ["building envelope", "plan of development", "pod plan",
                      "envelope plan", "development plan"],
        "min_matches": 1,
        "extensions": [".pdf", ".jpg", ".jpeg", ".png"],
    }),
    # Covenant Design Guidelines
    ("covenant_guidelines", {
        "keywords": ["design guidelines", "covenant guidelines", "architectural guidelines",
                      "design requirements", "building covenant"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Discount Approval
    ("discount_approval", {
        "keywords": ["discount approval", "discount form", "discount authorisation",
                      "discount authorization", "approved discount"],
        "min_matches": 1,
        "extensions": [".pdf", ".jpg", ".jpeg", ".png"],
    }),
    # Owner Supplied Items
    ("owner_supplied", {
        "keywords": ["owner supplied", "owner supply", "client supplied",
                      "owner provided items"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Modified Plan Approval
    ("modified_plan", {
        "keywords": ["modified plan", "plan modification", "plan amendment",
                      "modified plan approval"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Sales Accept Doc
    ("sales_accept", {
        "keywords": ["sales accept", "sales acceptance", "acceptance document"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
    # Covenant document (generic)
    ("covenant_doc", {
        "keywords": ["covenant", "building covenant", "restrictive covenant",
                      "community management statement"],
        "min_matches": 1,
        "extensions": [".pdf"],
    }),
]

# Spreadsheet fingerprints (for PSE Excel)
SPREADSHEET_EXTS = {".xlsx", ".xlsm", ".xls", ".csv", ".xlsb"}


def classify_by_content(file_info: dict, text_content: str, spreadsheet_info: dict | None) -> tuple[str | None, float]:
    """
    3-tier classification:
    1. Text content (strong, free) — works for text-based PDFs
    2. Filename hints (medium, free) — catches scanned PDFs and images with descriptive names
    3. Returns (None, 0.0) to trigger vision fallback for truly ambiguous files
    """
    ext = Path(file_info["name"]).suffix.lower()
    text_lower = text_content.lower() if text_content else ""
    fn_lower = file_info["name"].lower()

    # --- Tier 0: .docx files — extract text via python-docx if available ---
    if ext == ".docx":
        try:
            import docx as _docx
            doc = _docx.Document(file_info["full_path"])
            text_lower = " ".join(p.text for p in doc.paragraphs).lower()
        except Exception:
            text_lower = ""
        # Run filename hints first for .docx
        for patterns, doc_type in [
            (["intention to purchase", "itp"], "itp"),
            (["pse checklist", "checklist"], "pse_checklist"),
            (["provisional sales estimate", "pse"], "pse_doc"),
            (["disclosure plan", "survey plan"], "disclosure_plan"),
            (["promo", "acknowledgement", "acknowledgment"], "promo_ack"),
            (["red pen", "markup", "mark up"], "red_pen"),
        ]:
            if any(p in fn_lower for p in patterns):
                return doc_type, 0.75
        # Also check text content for PSE Checklist .docx
        if text_lower and ("pse documents" in text_lower or "provisional sales estimate documents" in text_lower
                           or "document checklist" in text_lower):
            return "pse_checklist", 0.85
        # Fall through to text-based classification below with extracted text

    # --- Tier 1a: Spreadsheet files → PSE Excel ---
    if ext in SPREADSHEET_EXTS:
        # Filename-first for xlsb: "PSE May 2026 v3 S26NLSP.xlsb" → PSE Excel
        # Also catches any spreadsheet with PSE in the filename
        if "pse" in fn_lower or "provisional" in fn_lower:
            return "pse_excel", 1.0
        if spreadsheet_info:
            # Check sheet names — PSE Excel has characteristic sheet names
            sheets_lower = [s.lower() for s in spreadsheet_info.get("sheets", [])]
            pse_sheet_signals = ["pse", "nhp", "gld pse", "gld nhp", "pse-nhp", "pse cover", "nhp cover"]
            if any(any(sig in s for sig in pse_sheet_signals) for s in sheets_lower):
                return "pse_excel", 1.0
            sample = spreadsheet_info.get("sample_text", "").lower()
            if any(kw in sample for kw in ["pse", "provisional", "price", "estimate", "inclusions",
                                            "base price", "total", "contract"]):
                return "pse_excel", 1.0
        # Any spreadsheet in a PSE submission zip is almost certainly the PSE Excel
        return "pse_excel", 0.8

    # --- Tier 1.5: Filename wins for Disclosure Plan to prevent acoustic mis-tag ---
    # A file named "disclosure plan" or "survey plan" should NEVER be tagged acoustic
    # even if it contains the word "acoustic" in passing.
    disclosure_fn_signals = ["disclosure plan", "survey plan", "plan of sub", "plan of survey"]
    if ext == ".pdf" and any(s in fn_lower for s in disclosure_fn_signals):
        return "disclosure_plan", 0.9

    # --- Tier 1b: PDF text-based classification ---
    if ext == ".pdf" and text_lower:
        # Priority override: PSE Checklist contains keywords for MANY doc types
        # (it lists them all). If the filename says "checklist", trust that.
        # Also catch "Site Visit Checklist" which is a different form.
        _checklist_text_signals = ["pse checklist", "document checklist", "1.0 pse document naming",
                                    "pse documents list", "document list", "document naming"]
        if "checklist" in fn_lower and any(s in text_lower for s in _checklist_text_signals):
            return "pse_checklist", 1.0
        # Site Visit Checklist — trust the filename
        if "site visit" in fn_lower and "checklist" in fn_lower:
            return "pse_checklist", 0.85

        # Priority override: ITP — "INTENTION TO PURCHASE" as a heading is definitive.
        # The ITP Acreage version mentions "provisional sales estimate" in the payment
        # section, which would otherwise cause pse_doc to win. ITP heading beats that.
        if "intention to purchase" in text_lower and (
            "purchaser" in text_lower or "purchaser details" in text_lower
        ):
            return "itp", 0.95

        # Priority override: Pool Form — POOL DECLARATION heading is definitive.
        # Pool forms contain "yes/no" and "completed" which otherwise match pse_checklist.
        if "pool declaration" in text_lower or (
            "swimming pool" in text_lower and "declaration" in text_lower
        ):
            return "pool_form", 0.95

        # Priority override: Advantage / Promo Ack — filename + content.
        # These contain "advantage" and "acknowledgement" which is unambiguous.
        if ("advantage" in fn_lower or "promo" in fn_lower or "acknowledgement" in fn_lower) and (
            "advantage" in text_lower or "super saver" in text_lower or "promotional" in text_lower
        ):
            return "promo_ack", 0.95

        # Priority override: Site Visit Checklist — not a GeoSite plan.
        if "site visit checklist" in text_lower or (
            "site visit" in text_lower and "checklist" in text_lower
        ):
            return "pse_checklist", 0.85

        best_type = None
        best_score = 0
        for doc_type, fingerprint in CONTENT_FINGERPRINTS:
            if ext not in fingerprint["extensions"]:
                continue
            matches = sum(1 for kw in fingerprint["keywords"] if kw in text_lower)
            if matches >= fingerprint["min_matches"] and matches > best_score:
                best_score = matches
                best_type = doc_type
        if best_type:
            confidence = min(1.0, 0.5 + best_score * 0.15)
            return best_type, confidence

    # --- Tier 1c: Known non-submission file exclusions ---
    # These files are sometimes included in zips but are NOT submission documents.
    # Return None so they end up in 'unclassified' rather than being misidentified.
    EXCLUDE_FILENAME_SIGNALS = [
        "8_step", "8 step", "process_timeline", "process timeline",
        "matters_of_interest", "matters of interest",
        "vegetation_management", "vegetation management",
        "generalmap", "general map",
        "zoning map", "council map",
        "ausmar_advantage",  # marketing doc - but Advantage Ack is classified by content above
    ]
    if any(s in fn_lower for s in EXCLUDE_FILENAME_SIGNALS):
        return None, 0.0  # Leave unclassified

    # --- Tier 2: Filename hints (for scanned PDFs and images) ---
    # These patterns are based on real AUSMAR consultant naming conventions.
    # Confidence is 0.75 — strong enough to use, but vision can override.
    FILENAME_HINTS = [
        # ITP
        (["intention to purchase", "itp form", "itp -", "- itp"], "itp"),
        # PSE Doc — also match bare 'pse' filename (scanned PSE PDFs often named PSE.pdf)
        (["provisional sales estimate", "pse doc", "pse -", "- pse", "pse signed", "pse.pdf"], "pse_doc"),
        # GeoSite
        (["geosite", "geo site", "geo plan", "geosite plan"], "geosite"),
        # Red Pen
        (["red pen", "redpen", "mark up", "markup", "red mark"], "red_pen"),
        # Pool Form
        (["swimming pool", "pool form"], "pool_form"),
        # Promo Ack
        (["promo", "client acknowledgement", "client acknowledgment"], "promo_ack"),
        # Disclosure Plan
        (["disclosure plan", "survey plan", "plan of sub"], "disclosure_plan"),
        # Deposit Receipt
        (["deposit remit", "deposit receipt", "remittance", "receipt"], "deposit_receipt"),
        # Drivers Licence
        (["drivers licence", "driver licence", "drivers license", "driver license", "dl -", "- dl "], "drivers_licence"),
        # POD / Building Envelope
        (["pod", "building envelope", "envelope plan"], "pod_envelope"),
        # Covenant Application
        (["covenant application", "covenant form", "covenant app"], "covenant_application"),
        # Sites with Fall
        (["sites with fall", "fall acknowledgment", "fall ack"], "fall_ack"),
        # Compaction Report
        (["compaction report", "compaction"], "compaction_report"),
        # Soil Report
        (["soil report", "soil test", "site classification"], "soil_report"),
        # Acoustic Report
        (["acoustic", "noise report"], "acoustic_report"),
        # BAL Report
        (["bal report", "bushfire"], "bal_report"),
        # Contour Survey
        (["contour survey", "contour plan", "topographic"], "contour_survey"),
        # Covenant Guidelines
        (["design guidelines", "covenant guidelines", "covenant doc"], "covenant_guidelines"),
        # PSE Checklist
        (["pse checklist", "document checklist", "pse documents list"], "pse_checklist"),
        # Discount Approval
        (["discount approval", "discount form"], "discount_approval"),
        # Owner Supplied
        (["owner supplied", "owner supply"], "owner_supplied"),
        # Modified Plan
        (["modified plan", "plan modification"], "modified_plan"),
    ]
    for patterns, doc_type in FILENAME_HINTS:
        if any(p in fn_lower for p in patterns):
            return doc_type, 0.75

    # Special case: bare filename is exactly 'pse' (e.g. PSE.pdf, pse.pdf)
    # These are scanned PSE Doc PDFs with no text content.
    fn_stem = Path(file_info["name"]).stem.lower().strip()
    if fn_stem == "pse" and ext == ".pdf":
        return "pse_doc", 0.75

    # --- Image files with no filename hint → vision ---
    if ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"):
        return None, 0.0  # Will be classified by LLM vision

    return None, 0.0


def classify_by_vision(file_info: dict) -> tuple[str | None, float]:
    """Use LLM vision to classify an image or scanned PDF. Returns (doc_type_key, confidence)."""
    ext = Path(file_info["name"]).suffix.lower()
    full_path = file_info["full_path"]

    try:
        if ext == ".pdf":
            pages_b64 = pdf_all_pages_to_base64(full_path, dpi=100, max_pages=1)
            if not pages_b64:
                return None, 0.0
            image_b64_list = pages_b64
        elif ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"):
            b64 = image_to_base64(full_path)
            if not b64:
                return None, 0.0
            image_b64_list = [b64]
        else:
            return None, 0.0

        doc_type_list = ", ".join(f'"{k}": {v}' for k, v in DOC_TYPES.items())

        system_prompt = f"""You are a document classifier for AUSMAR home building submissions.
Classify this document into ONE of these types:
{doc_type_list}

If it's a photo of a person's ID card or driver's licence, classify as "drivers_licence".
If it's a site plan with lot dimensions, setbacks, and house position, classify as "geosite".
If it's floor plans with red/coloured markings showing changes, classify as "red_pen".
If it's a building envelope or POD plan, classify as "pod_envelope".
If it's a deposit/payment receipt, classify as "deposit_receipt".
If it's a contour/topographic survey, classify as "contour_survey".
If it's a swimming pool declaration form, classify as "pool_form".
If it's an AUSMAR process timeline, step guide, or marketing document, classify as "unknown".
If it's a Matters of Interest report, vegetation management report, or council zoning map, classify as "unknown".
If it doesn't match any type, respond with "unknown".

Respond with ONLY a JSON object: {{"doc_type": "type_key", "confidence": 0.0-1.0, "reason": "brief reason"}}"""

        raw = call_vision_model(system_prompt, "Classify this document.", image_b64_list)
        result = parse_json_from_llm(raw)
        doc_type = result.get("doc_type", "unknown")
        confidence = float(result.get("confidence", 0.5))

        if doc_type and doc_type != "unknown" and doc_type in DOC_TYPES:
            return doc_type, confidence
        return None, 0.0

    except Exception as e:
        print(f"[WARN] classify_by_vision failed for {file_info['name']}: {e}")
        return None, 0.0


def classify_all_files(files: list[dict], progress_callback=None) -> dict:
    """
    Classify all files by content. Returns:
    {
        "file_map": {doc_type_key: file_info_with_classification},
        "unclassified": [file_info_list],
        "classifications": {filename: {"doc_type": ..., "confidence": ..., "method": ...}},
    }
    """
    file_map = {}  # doc_type -> file_info
    unclassified = []
    classifications = {}

    # Phase 1: Text-based classification (fast, free)
    for f in files:
        # Store original name before any renaming happens later
        f["original_name"] = f["name"]
        ext = Path(f["name"]).suffix.lower()
        text_content = ""
        spreadsheet_info = None

        if ext == ".pdf":
            text_content = extract_pdf_text(f["full_path"])
        elif ext in SPREADSHEET_EXTS:
            spreadsheet_info = extract_spreadsheet_info(f["full_path"])
            text_content = spreadsheet_info.get("sample_text", "") if spreadsheet_info else ""
        elif ext == ".docx":
            try:
                import docx as _docx
                _doc = _docx.Document(f["full_path"])
                text_content = " ".join(p.text for p in _doc.paragraphs)
            except Exception:
                text_content = ""

        doc_type, confidence = classify_by_content(f, text_content, spreadsheet_info)

        if doc_type and confidence >= 0.5:
            # If this doc_type already has a file, keep the higher confidence one
            if doc_type in file_map:
                existing = classifications.get(file_map[doc_type]["name"], {})
                if confidence > existing.get("confidence", 0):
                    # Demote existing to unclassified (or red_pen_extra for red pen)
                    if doc_type == "red_pen" and "red_pen_extra" not in file_map:
                        file_map["red_pen_extra"] = file_map[doc_type]
                    else:
                        unclassified.append(file_map[doc_type])
                    file_map[doc_type] = f
                    classifications[f["name"]] = {
                        "doc_type": doc_type, "confidence": confidence, "method": "text",
                        "extracted_text_preview": text_content[:200] if text_content else "",
                    }
                else:
                    # For red_pen duplicates, store as red_pen_extra instead of discarding
                    if doc_type == "red_pen" and "red_pen_extra" not in file_map:
                        file_map["red_pen_extra"] = f
                        classifications[f["name"]] = {
                            "doc_type": "red_pen_extra", "confidence": confidence, "method": "text",
                            "note": "Additional Red Pen file — will be merged into vision analysis",
                        }
                    else:
                        unclassified.append(f)
                        classifications[f["name"]] = {
                            "doc_type": doc_type, "confidence": confidence, "method": "text",
                            "note": f"Duplicate — {file_map[doc_type]['name']} already classified as {doc_type}",
                        }
            else:
                file_map[doc_type] = f
                classifications[f["name"]] = {
                    "doc_type": doc_type, "confidence": confidence, "method": "text",
                    "extracted_text_preview": text_content[:200] if text_content else "",
                }
        else:
            unclassified.append(f)
            classifications[f["name"]] = {
                "doc_type": None, "confidence": 0, "method": "text",
                "note": "No text match — queued for vision classification",
            }

    # Phase 2: Vision-based classification for unclassified files (costs API calls)
    still_unclassified = []
    for f in unclassified:
        ext = Path(f["name"]).suffix.lower()
        if ext not in (".pdf", ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".docx"):
            still_unclassified.append(f)
            continue

        doc_type, confidence = classify_by_vision(f)
        if doc_type and confidence >= 0.4:
            if doc_type in file_map:
                existing = classifications.get(file_map[doc_type]["name"], {})
                if confidence > existing.get("confidence", 0):
                    still_unclassified.append(file_map[doc_type])
                    file_map[doc_type] = f
                    classifications[f["name"]] = {
                        "doc_type": doc_type, "confidence": confidence, "method": "vision",
                    }
                else:
                    still_unclassified.append(f)
            else:
                file_map[doc_type] = f
                classifications[f["name"]] = {
                    "doc_type": doc_type, "confidence": confidence, "method": "vision",
                }
        else:
            still_unclassified.append(f)
            classifications[f["name"]] = classifications.get(f["name"], {})
            classifications[f["name"]]["vision_result"] = "unclassified"

    return {
        "file_map": file_map,
        "unclassified": still_unclassified,
        "classifications": classifications,
    }


# ---------------------------------------------------------------------------
# Check 1: File Structure (extract, flatten, remove junk)
# ---------------------------------------------------------------------------
def check_file_structure(extract_dir: str, zip_name: str) -> dict:
    issues = []
    warnings = []
    corrections = []
    files = []

    for root, dirs, filenames in os.walk(extract_dir):
        rel_root = os.path.relpath(root, extract_dir)
        for fn in filenames:
            if fn.startswith(".") or "__MACOSX" in root:
                continue
            rel_path = os.path.join(rel_root, fn) if rel_root != "." else fn
            full_path = os.path.join(root, fn)
            files.append({"name": fn, "rel_path": rel_path, "full_path": full_path})

    # Flatten subfolders
    has_subfolders = False
    for f in files:
        if os.path.dirname(f["rel_path"]) and os.path.dirname(f["rel_path"]) != ".":
            has_subfolders = True
            new_path = os.path.join(extract_dir, f["name"])
            if not os.path.exists(new_path):
                shutil.move(f["full_path"], new_path)
                corrections.append(f"Flattened: moved '{f['rel_path']}' to zip root")
                f["full_path"] = new_path
                f["rel_path"] = f["name"]

    if has_subfolders:
        corrections.append("Removed subfolder structure — all files now flat in zip root")
        for root, dirs, _ in os.walk(extract_dir, topdown=False):
            for d in dirs:
                dp = os.path.join(root, d)
                try:
                    os.rmdir(dp)
                except Exception:
                    pass

    # Extract deal code from zip name (no warning — consultants name zips inconsistently)
    zip_stem = Path(zip_name).stem
    # Try to extract deal code pattern from the zip name even if it has extra words
    deal_code_match = re.search(r'[A-Za-z]\d{2}[A-Za-z]{2,6}', zip_stem)
    if deal_code_match:
        zip_stem = deal_code_match.group(0)

    # Remove junk files
    cleaned_files = []
    seen_names = set()

    for f in files:
        fn = f["name"]
        fn_lower = fn.lower()

        is_junk = any(pat in fn_lower for pat in JUNK_PATTERNS)
        if fn_lower.endswith(".msg"):
            is_junk = True
        if is_junk:
            try:
                os.remove(f["full_path"])
            except Exception:
                pass
            corrections.append(f"Removed junk file: '{fn}'")
            continue

        if f["name"].lower() in seen_names:
            issues.append(f"Duplicate file: '{f['name']}'")
        seen_names.add(f["name"].lower())
        cleaned_files.append(f)

    return {
        "files": cleaned_files,
        "issues": issues,
        "warnings": warnings,
        "corrections": corrections,
        "zip_stem": zip_stem,
    }


# ---------------------------------------------------------------------------
# Check 2: Document Completeness (content-based)
# ---------------------------------------------------------------------------
def check_document_completeness(file_map: dict, unclassified: list) -> dict:
    """Check completeness using the content-classified file_map, NOT filenames."""
    issues = []
    warnings = []
    found = {}
    found_conditional = []

    # Core required
    for doc_key in CORE_REQUIRED:
        canonical = DOC_TYPES[doc_key]
        if doc_key in file_map:
            found[canonical] = file_map[doc_key]["name"]
        else:
            issues.append(f"Missing: {canonical}")

    # Soft required
    for doc_key in SOFT_REQUIRED:
        canonical = DOC_TYPES[doc_key]
        if doc_key in file_map:
            found[canonical] = file_map[doc_key]["name"]
        else:
            warnings.append(f"Missing: {canonical} — should be included per 1.0 PSE Document Naming")

    # Conditional documents — note which are present
    conditional_keys = [
        "pod_envelope", "compaction_report", "covenant_guidelines", "disclosure_plan",
        "covenant_application", "pool_form", "discount_approval", "owner_supplied",
        "modified_plan", "fall_ack", "acoustic_report", "bal_report",
        "contour_survey", "sales_accept", "soil_report", "covenant_doc",
    ]
    for doc_key in conditional_keys:
        if doc_key in file_map:
            found_conditional.append(DOC_TYPES[doc_key])

    # Expected conditionals
    for doc_key in CONDITIONAL_EXPECTED:
        canonical = DOC_TYPES[doc_key]
        if doc_key not in file_map:
            warnings.append(f"Missing: {canonical} — required per 1.0 naming unless explained")

    # Drivers licence format check
    if "drivers_licence" in file_map:
        dl_ext = Path(file_map["drivers_licence"]["name"]).suffix.lower()
        if dl_ext in LICENCE_BAD_EXTS:
            issues.append(f"Drivers Licence in {dl_ext} format — needs .jpg, .png, or .pdf")

    # Note unclassified files
    if unclassified:
        unc_names = [f["name"] for f in unclassified]
        warnings.append(f"Unclassified files (could not determine document type): {', '.join(unc_names)}")

    return {
        "found": found,
        "issues": issues,
        "warnings": warnings,
        "conditional_found": found_conditional,
    }


# ---------------------------------------------------------------------------
# Rename files to standard names based on classification
# ---------------------------------------------------------------------------
def rename_classified_files(extract_dir: str, file_map: dict, classifications: dict) -> list[str]:
    """Rename files to their canonical names. Returns list of correction descriptions."""
    corrections = []

    for doc_key, file_info in file_map.items():
        canonical_name = DOC_TYPES.get(doc_key)
        if not canonical_name:
            continue

        old_name = file_info["name"]
        ext = Path(old_name).suffix
        new_name = canonical_name + ext

        if new_name != old_name:
            old_path = file_info["full_path"]
            new_path = os.path.join(extract_dir, new_name)

            # Avoid overwriting
            if os.path.exists(new_path) and new_path != old_path:
                # Add a suffix to avoid collision
                stem = canonical_name
                counter = 2
                while os.path.exists(new_path):
                    new_name = f"{stem} ({counter}){ext}"
                    new_path = os.path.join(extract_dir, new_name)
                    counter += 1

            try:
                os.rename(old_path, new_path)
                corrections.append(f"Renamed '{old_name}' -> '{new_name}' (identified as {canonical_name} by content)")
                file_info["name"] = new_name
                file_info["full_path"] = new_path
                file_info["rel_path"] = new_name
            except Exception as e:
                print(f"[WARN] rename failed for {old_name}: {e}")

    return corrections


# ---------------------------------------------------------------------------
# Check 3: GeoSite Verification (Vision) — uses file_map
# ---------------------------------------------------------------------------
def check_geosite(file_map: dict, land_registered: bool | None = None) -> dict:
    """
    land_registered: True = land registered (full siting plan required).
                     False/None = unregistered or unknown (lot/locality plan is acceptable).
    """
    if "geosite" not in file_map:
        return {
            "issues": ["No GeoSite file found in submission — CRITICAL per 1.0 PSE Document Naming"],
            "warnings": [], "analysis": {}, "lot_dimensions": None,
        }

    geosite_file = file_map["geosite"]
    geosite_path = geosite_file["full_path"]
    ext = Path(geosite_path).suffix.lower()

    try:
        vision_warning = None
        if ext == ".pdf":
            pages_b64 = pdf_all_pages_to_base64(geosite_path, dpi=100, max_pages=3)
            if not pages_b64:
                vision_warning = "PDF vision analysis skipped — poppler unavailable or PDF conversion failed"
        elif ext in (".jpg", ".jpeg", ".png"):
            b64 = image_to_base64(geosite_path)
            pages_b64 = [b64] if b64 else []
            if not pages_b64:
                vision_warning = "Image conversion failed for GeoSite"
        else:
            return {
                "issues": [f"GeoSite in unsupported format: {ext}"],
                "warnings": [], "analysis": {}, "lot_dimensions": None,
            }

        if not pages_b64:
            return {
                "issues": [],
                "warnings": [vision_warning or "Could not convert GeoSite to images — vision analysis skipped"],
                "analysis": {}, "lot_dimensions": None,
            }

        fp_notes = _get_fp_notes("geosite_verification")

        system_prompt = f"""You are an AUSMAR QA reviewer analysing a GeoSite siting document.

A VALID GeoSite from geosite.com.au looks like this:
- Header: "Proposed siting of your new AUSMAR Home" (or similar AUSMAR siting header)
- Bottom-left: "© GeoSite IT Pty Ltd" watermark AND/OR a "Geo Plan ID" number
- Bottom-left: North arrow compass
- Right panel: Customer details (Customer, Site Address, Locality, Home Design, Estate, SP Number, Date)
- Right panel: Site stats box (Site Area, Site Coverage, Build Area, Ceiling Height)
- Main area: House floor plan positioned inside the lot boundary (lot shown in black outline, house in pink/red outline)
- Dimension annotations showing lot measurements and setbacks (numbers with "m" suffix)
- Customer signature lines (usually bottom-left, may be signed or blank)
- Consultant name and email (bottom-left area)

IMPORTANT CALIBRATION — only flag these as TRUE problems:
- is_geosite_tool = false ONLY if the document is clearly NOT a GeoSite (e.g. it's a hand-drawn sketch, a site visit checklist, or a contour survey with no house positioned on it). If you can see the AUSMAR siting header OR the GeoSite IT Pty Ltd watermark OR a Geo Plan ID, set this to true.
- house_sited_at_scale = false ONLY if there is literally no house floor plan visible on the lot. If a floor plan is positioned inside the lot boundary, set this to true.
- setback_dimensions_shown = false ONLY if there are NO numeric dimension annotations at all. If you can see any measurements (even just 2-3 numbers with "m"), set this to true.
- customer_signatures_present = false ONLY if the signature lines are clearly blank/empty. If there are any marks, initials, or signatures, set this to true.
- is_combined_with_contours = true ONLY if contour lines are overlaid on top of the GeoSite making the plan hard to read.

Extract data and report in JSON format:
{{
  "is_geosite_tool": true/false,
  "is_combined_with_contours": true/false,
  "house_sited_at_scale": true/false,
  "setback_dimensions_shown": true/false,
  "text_readable": true/false,
  "geo_plan_id_visible": true/false/null,
  "consultant_name": "name or null",
  "customer_signatures_present": true/false,
  "site_coverage_percent": number or null,
  "site_area_m2": number or null,
  "build_area_m2": number or null,
  "home_design": "plan name or null",
  "facade_name": "facade name or null",
  "estate_name": "name or null",
  "lot_number": "number or null",
  "street_address": "address or null",
  "sp_number": "value or null",
  "lot_width_m": number or null,
  "lot_length_m": number or null,
  "lot_area_m2": number or null,
  "side_setback_left_m": number or null,
  "side_setback_right_m": number or null,
  "front_setback_m": number or null,
  "rear_setback_m": number or null,
  "fall_across_site_mm": number or null,
  "is_battle_axe_lot": true/false/null,
  "concerns": ["list of concerns — only include if genuinely problematic"],
  "notes": "additional observations"
}}

Be conservative: only flag something as false/problematic if you are CERTAIN it is wrong. When in doubt, assume the document is valid.{fp_notes}"""

        raw = call_vision_model(
            system_prompt,
            "Analyse this GeoSite document. Extract all dimensions, check all required elements, and identify any concerns.",
            pages_b64,
        )
        analysis = parse_json_from_llm(raw)

        issues = []
        warnings = []

        # Context-aware: only require full siting plan if land IS registered.
        # If land is NOT registered (or unknown/None), a lot/locality plan is acceptable.
        require_full_siting = (land_registered is True)

        if analysis.get("is_geosite_tool") is False:
            if require_full_siting:
                warnings.append(
                    "Document may not be from geosite.com.au — verify it shows AUSMAR siting header "
                    "or GeoSite IT Pty Ltd watermark. If it is a valid GeoSite, ignore this warning."
                )
            else:
                warnings.append(
                    "Document appears to be a lot/locality plan — acceptable as land is not yet registered. "
                    "Once land registers, a full GeoSite siting plan will be required."
                )
        if analysis.get("is_combined_with_contours") is True:
            warnings.append(
                "GeoSite appears combined with contour data — Heath requires these to be separate documents. "
                "Contours overlaid on GeoSite make it unreadable (real issue from S26TLS review)."
            )
        if analysis.get("house_sited_at_scale") is False:
            if require_full_siting:
                issues.append("House not sited at scale on the lot — cannot verify fit")
            else:
                warnings.append(
                    "House not yet sited on lot plan — acceptable as land is not yet registered. "
                    "Full GeoSite siting required once land registers."
                )
        if analysis.get("setback_dimensions_shown") is False:
            if require_full_siting:
                issues.append(
                    "Setback dimensions not shown on GeoSite — MUST have all setbacks "
                    "(front, rear, left side, right side) for drafting team"
                )
            else:
                warnings.append(
                    "No setback dimensions on lot plan — acceptable as land is not yet registered. "
                    "Setbacks required on GeoSite once land registers."
                )
        if analysis.get("text_readable") is False:
            warnings.append("Text on GeoSite is overlapping or hard to read — may cause issues for drafting")
        if analysis.get("customer_signatures_present") is False:
            warnings.append("Customer signature(s) may be missing from GeoSite — required per 1.0 naming")

        # Detect regional rules for setback/coverage checks
        estate_for_check = analysis.get("estate_name") or ""
        addr_for_check = analysis.get("street_address") or ""
        reg_rules = detect_regional_rules(estate_for_check, addr_for_check)
        
        front_sb = analysis.get("front_setback_m")
        if front_sb is not None and isinstance(front_sb, (int, float)):
            min_front = (reg_rules or {}).get("front_setback_m", 3.0)
            if front_sb < min_front:
                warnings.append(
                    f"Front setback is {front_sb}m — minimum is {min_front}m "
                    f"({'per ' + reg_rules['name'] if reg_rules else 'verify with council'}). "
                    f"Verify against council requirements."
                )

        for side_key, side_label in [("side_setback_left_m", "Left"), ("side_setback_right_m", "Right")]:
            val = analysis.get(side_key)
            if val is not None and isinstance(val, (int, float)):
                min_side = (reg_rules or {}).get("side_setback_m", 0.75)
                if val < min_side:
                    # Demoted to WARNING — build-to-boundary lots are valid per covenant
                    warnings.append(
                        f"{side_label} side setback is {val}m — minimum is {min_side}m "
                        f"({'per ' + reg_rules['name'] if reg_rules else 'check council requirements'}). "
                        f"Verify against covenant — may be a build-to-boundary lot."
                    )

        coverage = analysis.get("site_coverage_percent")
        if coverage is not None and isinstance(coverage, (int, float)):
            max_cov = (reg_rules or {}).get("max_site_coverage_pct", 60)
            if coverage > max_cov:
                issues.append(
                    f"Site coverage {coverage}% exceeds {max_cov}% maximum "
                    f"({'per ' + reg_rules['name'] if reg_rules else 'check council requirements'})"
                )
            elif coverage > max_cov - 2:
                warnings.append(
                    f"Site coverage {coverage}% is very close to {max_cov}% maximum — zero margin. "
                    f"Verify with covenant."
                )

        if analysis.get("is_battle_axe_lot") is True:
            warnings.append(
                "Lot appears to be battle-axe configuration — may require LHDC "
                "(Livable Housing Design Standards) assessment. Verify with council."
            )

        fall_mm = analysis.get("fall_across_site_mm")
        if fall_mm is not None and isinstance(fall_mm, (int, float)):
            if fall_mm >= 500:
                warnings.append(
                    f"Site fall detected: {fall_mm}mm — Sites with Fall Acknowledgment (Signed) required. "
                    f"Max cut/fill before building manager approval is 1000mm."
                )
            if fall_mm >= 1000:
                issues.append(
                    f"Significant site fall: {fall_mm}mm — may require contour survey, "
                    f"retaining wall engineering, and council application"
                )

        # Home Design field: only warn if we couldn't extract it AND the plan name wasn't
        # found elsewhere — this was a FP when the LLM couldn't read small text
        if not analysis.get("home_design"):
            warnings.append(
                "Home Design field not visible on GeoSite — verify plan name is filled in. "
                "If it is filled in, the text may be too small for automated reading."
            )

        lot_dims = {
            "width": analysis.get("lot_width_m"),
            "length": analysis.get("lot_length_m"),
            "area": analysis.get("lot_area_m2") or analysis.get("site_area_m2"),
        }

        return {
            "issues": issues, "warnings": warnings,
            "analysis": analysis, "lot_dimensions": lot_dims,
        }

    except Exception as e:
        traceback.print_exc()
        return {
            "issues": [],
            "warnings": [f"GeoSite vision analysis skipped due to error: {str(e)}"],
            "analysis": {}, "lot_dimensions": None,
        }


# ---------------------------------------------------------------------------
# Check 4: Plan-to-Lot Fit
# ---------------------------------------------------------------------------
def check_plan_to_lot_fit(geosite_result: dict, file_map: dict) -> dict:
    issues = []
    warnings = []
    analysis = geosite_result.get("analysis", {})
    lot_dims = geosite_result.get("lot_dimensions", {})

    if not analysis or isinstance(analysis, str):
        return {
            "issues": ["Cannot verify plan-to-lot fit — GeoSite analysis unavailable"],
            "warnings": [], "plan_identified": None, "fit_result": "UNKNOWN",
        }

    home_design = analysis.get("home_design", "") or ""

    try:
        plans = db.get_all_plans()
    except Exception:
        plans = []
    matched_plan = None

    for p in plans:
        if p["name"].lower() in home_design.lower():
            matched_plan = p
            break

    if not matched_plan:
        for p in plans:
            words = p["name"].lower().split()
            if all(w in home_design.lower() for w in words):
                matched_plan = p
                break

    if not matched_plan:
        for p in plans:
            family = p["name"].lower().split()[0]
            if len(family) > 3 and family in home_design.lower():
                matched_plan = p
                break

    if not matched_plan:
        note = (
            f"Plan '{home_design}' is not in the current AUSMAR plan library "
            f"(Designer, Boutique, Acreage collections). This is likely a discontinued "
            f"or legacy plan — manual plan-to-lot fit check required by Heath."
            if home_design
            else "Could not identify plan name from GeoSite — manual plan-to-lot fit check required by Heath"
        )
        warnings.append(note)
        return {
            "issues": issues, "warnings": warnings,
            "plan_identified": home_design, "fit_result": "MANUAL CHECK REQUIRED",
            "plan_specs": None, "lot_dimensions": lot_dims,
        }

    lot_width = lot_dims.get("width") if lot_dims else None
    lot_length = lot_dims.get("length") if lot_dims else None
    fit_result = "PASS"

    if lot_width is not None and isinstance(lot_width, (int, float)):
        if lot_width < matched_plan["min_width"]:
            issues.append(
                f"CRITICAL: {matched_plan['name']} requires minimum {matched_plan['min_width']}m site width "
                f"but lot is only {lot_width}m wide. Plan DOES NOT FIT."
            )
            fit_result = "FAIL — PLAN TOO WIDE"
        elif lot_width < matched_plan["min_width"] + 0.5:
            warnings.append(
                f"Tight fit: {matched_plan['name']} needs {matched_plan['min_width']}m width, "
                f"lot is {lot_width}m. Only {lot_width - matched_plan['min_width']:.1f}m margin."
            )
            fit_result = "TIGHT FIT"
    else:
        warnings.append("Could not extract lot width from GeoSite — manual verification needed")

    if lot_length is not None and isinstance(lot_length, (int, float)):
        if lot_length < matched_plan["min_length"]:
            issues.append(
                f"CRITICAL: {matched_plan['name']} requires minimum {matched_plan['min_length']}m depth "
                f"but lot is only {lot_length}m. Plan DOES NOT FIT."
            )
            if fit_result == "PASS":
                fit_result = "FAIL — LOT TOO SHORT"

    build_area = analysis.get("build_area_m2")
    if build_area and matched_plan.get("total_area"):
        try:
            ba = float(str(build_area).replace("m2", "").replace("m²", "").strip())
            if ba > matched_plan["total_area"] * 1.15:
                warnings.append(
                    f"Build area ({ba}m²) significantly exceeds {matched_plan['name']} "
                    f"standard area ({matched_plan['total_area']}m²) — verify modifications"
                )
        except (ValueError, TypeError):
            pass

    return {
        "issues": issues, "warnings": warnings,
        "plan_identified": matched_plan["name"], "plan_specs": matched_plan,
        "lot_dimensions": lot_dims, "fit_result": fit_result,
    }


# ---------------------------------------------------------------------------
# Check 5: Red Pen Markup (Vision) — uses file_map
# ---------------------------------------------------------------------------
def check_red_pen(file_map: dict, deposit_type: str) -> dict:
    if "red_pen" not in file_map:
        if deposit_type == "NHP":
            return {
                "issues": ["Missing Red Pen Markup — required for NHP ($2,500) submissions per 1.0 naming"],
                "warnings": [], "analysis": {},
            }
        else:
            return {
                "issues": [],
                "warnings": ["No Red Pen Markup found — acceptable for STC ($4,000) if no structural changes"],
                "analysis": {},
            }

    redpen_file = file_map["red_pen"]
    redpen_path = redpen_file["full_path"]
    ext = Path(redpen_path).suffix.lower()

    try:
        vision_warning = None
        if ext == ".pdf":
            pages_b64 = pdf_all_pages_to_base64(redpen_path, dpi=100, max_pages=5)
            if not pages_b64:
                vision_warning = "PDF vision analysis skipped — poppler unavailable or PDF conversion failed"
        elif ext in (".jpg", ".jpeg", ".png"):
            b64 = image_to_base64(redpen_path)
            pages_b64 = [b64] if b64 else []
            if not pages_b64:
                vision_warning = "Image conversion failed for Red Pen"
        else:
            return {
                "issues": [f"Red Pen in unsupported format: {ext}"],
                "warnings": [], "analysis": {},
            }

        # Merge pages from red_pen_extra (second red pen file) into the analysis
        if "red_pen_extra" in file_map:
            extra_path = file_map["red_pen_extra"]["full_path"]
            extra_ext = Path(extra_path).suffix.lower()
            try:
                if extra_ext == ".pdf":
                    extra_pages = pdf_all_pages_to_base64(extra_path, dpi=100, max_pages=5)
                    if extra_pages:
                        pages_b64 = (pages_b64 or []) + extra_pages
                        print(f"[INFO] Merged {len(extra_pages)} pages from red_pen_extra into Red Pen analysis")
                elif extra_ext in (".jpg", ".jpeg", ".png"):
                    b64 = image_to_base64(extra_path)
                    if b64:
                        pages_b64 = (pages_b64 or []) + [b64]
            except Exception as e:
                print(f"[WARN] Failed to merge red_pen_extra pages: {e}")

        if not pages_b64:
            return {
                "issues": [],
                "warnings": [vision_warning or "Could not convert Red Pen to images — vision analysis skipped"],
                "analysis": {},
            }

        fp_notes = _get_fp_notes("red_pen_markup")

        # ---- Filename-based plan type pre-detection (runs BEFORE vision call) ----
        # Extract confirmed plan types from filenames so we can:
        # (a) tell the LLM what's in the document, and
        # (b) force-override any LLM misses after parsing.
        FILENAME_PLAN_TOKENS = [
            ("floor_plan",      ["floor plan", " floor,", " floor ", "floor.pdf", "flr plan"]),
            ("elevations",      ["elev", "elevation"]),
            ("electrical",      ["elec", "electrical", "elect"]),
            ("floor_coverings", ["f cov", "fcov", "floor cov", "floor covering", "flr cov"]),
            ("concrete",        ["concrete", "slab"]),
            ("site_plan",       ["site plan", " site,", " site "]),
            ("kitchen",         ["kit,", " kit ", "kitchen"]),
            ("bathroom",        ["ensuite", "bath", "bathroom"]),
        ]
        rp_combined_name = " ".join([
            redpen_file.get("name", "").lower(),
            file_map.get("red_pen_extra", {}).get("name", "").lower(),
        ])
        filename_confirmed_types = set()
        for plan_key, tokens in FILENAME_PLAN_TOKENS:
            for tok in tokens:
                if tok in rp_combined_name:
                    filename_confirmed_types.add(plan_key)
                    break
        if filename_confirmed_types:
            print(f"[INFO] Red pen filename pre-confirms plan types: {filename_confirmed_types}")

        deposit_desc = (
            "NHP ($2,500) — red pen markups ARE required, must be RED on AUSMAR base plan with dimensions"
            if deposit_type == "NHP"
            else "STC ($4,000) — clean plans expected, no structural changes"
        )

        system_prompt = f"""You are an AUSMAR QA reviewer analysing Red Pen Markup documents.
This is a {deposit_desc} submission.

WHAT A VALID AUSMAR RED PEN MARKUP LOOKS LIKE:
- It is a multi-page PDF with AUSMAR standard floor plans as the base
- Changes are marked with COLORED highlights (green, yellow, blue, pink, or red circles/highlights)
- The term "red pen" is a legacy name — modern AUSMAR markups use COLORED HIGHLIGHTS, not literal red ink
- Numbered reference circles (e.g., 3.0, 3.2.a) mark each change area
- Customer initials and date appear on each page (e.g., "TB LB 15/02/2026")
- Multiple pages covering: floor plan, elevations, electrical/floor coverings, concrete/slab
- The base plan shows AUSMAR branding, window schedules, floor area tables
- Changes may include raked ceilings, room modifications, fixture selections

CALIBRATION — BE CONSERVATIVE:
- is_red_colour: Set TRUE if markups use ANY coloured highlights (green, yellow, blue, pink, red). Only set FALSE if there are literally NO coloured markings at all.
- is_on_ausmar_base_plan: Set TRUE if the base plan has AUSMAR branding, standard room layouts, window schedules, or floor area tables. Only set FALSE if the plan is clearly from a different builder or hand-drawn.
- has_dimensions_on_changes: Set TRUE if you can see ANY dimension annotations on changed areas. Only set FALSE if changes are shown but have zero dimensions.
- customer_signed: Set TRUE if you can see ANY initials, signatures, or date markings. Only set FALSE if signature areas are clearly blank.
- For plan_types_covered: A page showing a floor plan with electrical symbols AND floor covering notes counts as covering BOTH electrical and floor_coverings.

Analyse and report in JSON:
{{
  "markup_colour": "describe the actual colours used (e.g., green/yellow/blue highlights)",
  "is_red_colour": true/false,
  "is_on_ausmar_base_plan": true/false,
  "has_dimensions_on_changes": true/false,
  "customer_signed": true/false,
  "plan_types_covered": ["list: floor_plan, elevations, electrical, floor_coverings, concrete"],
  "missing_plan_types": ["list of required types not found — empty list if all covered"],
  "structural_changes_shown": true/false,
  "changes_description": "brief description of what changes are shown",
  "reference_numbers_found": ["list of reference numbers like 3.0, 3.2.a found on the markup"],
  "tags_reference_pse_sections": true/false/null,
  "hebel_changeover_noted": true/false,
  "width_reductions_across_plan": true/false,
  "facade_changes_shown": true/false,
  "window_deletions_noted": true/false/null,
  "concerns": ["only include genuinely problematic items"],
  "notes": ""
}}

IMPORTANT: Be CONSERVATIVE. Only flag issues you are CERTAIN about. When in doubt, assume the markup is valid. Colored highlights (green, yellow, blue, pink) ARE valid markup colours for AUSMAR red pen documents.{fp_notes}
{f'FILENAME CONFIRMS these plan types are present in this document: {sorted(filename_confirmed_types)}. Include ALL of these in plan_types_covered and do NOT list them in missing_plan_types.' if filename_confirmed_types else ''}"""

        raw = call_vision_model(system_prompt, "Analyse these Red Pen Markup pages.", pages_b64)
        analysis = parse_json_from_llm(raw)

        issues = []
        warnings = []

        if deposit_type == "NHP":
            # CRITICAL checks — all three block acceptance:
            # 1. Must be on AUSMAR base plan
            # 2. Must be in colour
            # 3. Must include full set of plan types (to prove changes or no changes)
            # 4. Must be signed/initialled by customer
            if analysis.get("is_on_ausmar_base_plan") is False:
                issues.append(
                    "Markups NOT on standard AUSMAR base plan — "
                    "must overlay AUSMAR plan, not consultant's own program "
                    "(real rejection reason from S26TLS)"
                )
            if analysis.get("is_red_colour") is False:
                colour = analysis.get("markup_colour", "unknown")
                issues.append(
                    f"Red Pen markups are not in colour (appears {colour}) — "
                    f"all changes must be highlighted in colour so drafting can identify them"
                )
            if analysis.get("customer_signed") is False:
                issues.append(
                    "Red Pen is not signed/initialled by the customer — "
                    "customer sign-off is required on all Red Pen pages"
                )
            # DEFINITIVE filename override: runs unconditionally after LLM response.
            # filename_confirmed_types was computed before the vision call.
            # Normalise missing_plan_types — LLM may return None, string, or use spaces instead of underscores
            raw_missing = analysis.get("missing_plan_types") or []
            if isinstance(raw_missing, str):
                raw_missing = [x.strip() for x in raw_missing.replace(",", " ").split() if x.strip()]
            # Normalise keys: "floor coverings" → "floor_coverings" etc.
            def _norm_key(k):
                return k.lower().replace(" ", "_").replace("-", "_")
            missing = [_norm_key(m) for m in raw_missing if m]
            if filename_confirmed_types:
                before = list(missing)
                missing = [m for m in missing if m not in filename_confirmed_types]
                print(f"[INFO] Red pen filename override: confirmed={filename_confirmed_types}, before={before}, after={missing}")
                # NUCLEAR: delete missing_plan_types from analysis entirely so verdict LLM
                # never sees them as missing. Also update plan_types_covered.
                analysis.pop("missing_plan_types", None)
                analysis["plan_types_confirmed_by_filename"] = sorted(filename_confirmed_types)
                existing_covered = analysis.get("plan_types_covered") or []
                if isinstance(existing_covered, str):
                    existing_covered = [existing_covered]
                analysis["plan_types_covered"] = sorted(set([_norm_key(x) for x in existing_covered]) | filename_confirmed_types)
                # Also remove any 'concerns' entries that mention the confirmed types
                concerns = analysis.get("concerns", [])
                if isinstance(concerns, list):
                    confirmed_keywords = [t.replace("_", " ") for t in filename_confirmed_types]
                    analysis["concerns"] = [
                        c for c in concerns
                        if not any(kw in str(c).lower() for kw in confirmed_keywords)
                    ]
            if missing and len(missing) > 0:
                issues.append(
                    f"Red Pen is missing required plan types: {', '.join(missing)}. "
                    f"Full set required (Floor Plan, Elevations, Electrical, Floor Coverings, Concrete) "
                    f"to confirm whether changes exist or not on each page."
                )
            if analysis.get("has_dimensions_on_changes") is False:
                warnings.append(
                    "Some changed areas on Red Pen may not have dimensions — "
                    "verify all structural changes are dimensioned"
                )

        elif deposit_type == "STC":
            if analysis.get("structural_changes_shown") is True:
                issues.append("Structural changes on STC submission — STC should have clean plans")

        if analysis.get("hebel_changeover_noted"):
            warnings.append(
                "Hebel changeover noted — may be a setback compliance workaround. "
                "Heath should verify this is legitimate."
            )
        if analysis.get("width_reductions_across_plan"):
            warnings.append(
                "Width reductions across entire plan — possible force-fit to lot. "
                "Verify plan-to-lot fit carefully."
            )
        # Only warn about tag/PSE section mismatch if the model actually found reference numbers
        # on the markup — if no ref numbers found, this check is meaningless
        ref_nums = analysis.get("reference_numbers_found", [])
        if analysis.get("tags_reference_pse_sections") is False and ref_nums:
            warnings.append(
                "Red pen reference numbers found but may not match PSE section write-ups "
                "(e.g. tags like 3.2.a should have corresponding PSE sections) — verify with Heath"
            )

        return {"issues": issues, "warnings": warnings, "analysis": analysis}

    except Exception as e:
        traceback.print_exc()
        return {
            "issues": [],
            "warnings": [f"Red Pen vision analysis skipped due to error: {str(e)}"],
            "analysis": {},
        }


# ---------------------------------------------------------------------------
# Check 6: PSE Excel Analysis — uses file_map
# ---------------------------------------------------------------------------
def check_pse_excel(file_map: dict, geosite_analysis: dict, results: dict = None) -> dict:
    """Check PSE Excel: extract totals, compare with signed PDF, check pricing period."""
    if "pse_excel" not in file_map:
        return {
            "issues": ["Missing: PSE Excel — required per 1.0 PSE Document Naming"],
            "warnings": [], "analysis": {},
        }

    issues = []
    warnings = []
    analysis = {}
    print(f"[DEBUG] check_pse_excel: pse_excel found, pse_doc in file_map: {'pse_doc' in file_map}")

    facade_name = geosite_analysis.get("facade_name", "") or ""
    if facade_name:
        analysis["geosite_facade"] = facade_name

    # Gas cooktops: NOT flagged. Stockland ruling overrides covenant restrictions.

    # === Extract totals from PSE Excel ===
    excel_path = file_map["pse_excel"]["full_path"]
    # Use original filename for pricing period extraction (before renaming)
    original_excel_name = file_map["pse_excel"].get("original_name", "")
    excel_totals = extract_pse_excel_totals(excel_path, original_name=original_excel_name)
    analysis["excel_totals"] = excel_totals

    if excel_totals.get("proposed_price"):
        analysis["excel_proposed_price"] = excel_totals["proposed_price"]
    print(f"[DEBUG] check_pse_excel: proposed_price={excel_totals.get('proposed_price')}, pricing_period={excel_totals.get('pricing_period')}")

    # === CRITICAL CHECK: Compare PSE Excel total vs PSE Signed PDF total ===
    if excel_totals.get("proposed_price") and "pse_doc" in file_map:
        try:
            pse_path = file_map["pse_doc"]["full_path"]
            # The totals (TOTAL DEBITS, TOTAL CREDITS, PROPOSED NEW HOME PRICE)
            # are typically around page 11 of a 15-page PSE (60-80% mark).
            # Convert page-by-page to save memory on 512MB container.
            from pypdf import PdfReader as _PdfReader
            try:
                _reader = _PdfReader(pse_path)
                total_pages = len(_reader.pages)
            except Exception:
                total_pages = 12  # Assume ~12 pages if we can't read
            
            # Target pages at 60-80% mark
            start_page = max(0, int(total_pages * 0.6) - 1)  # 0-indexed
            end_page = min(total_pages - 1, int(total_pages * 0.85))
            pages_to_check = []
            for pg in range(start_page, end_page + 1):
                b64 = pdf_page_to_base64(pse_path, page_num=pg, dpi=100)
                if b64:
                    pages_to_check.append(b64)
            
            print(f"[DEBUG] PSE price check: total_pages={total_pages}, checking pages {start_page+1}-{end_page+1} ({len(pages_to_check)} converted)")
            if pages_to_check:
                raw = call_vision_model(
                    "You are extracting pricing data from an AUSMAR Provisional Sales Estimate (PSE) document. "
                    "Find the TOTAL DEBITS, TOTAL CREDITS, and PROPOSED NEW HOME PRICE values. "
                    "These are typically in a summary table near the end of the document. "
                    "Return ONLY a JSON object: "
                    '{"total_debits": 123456, "total_credits": 12345, "proposed_new_home_price": 123456}. '
                    "Return the numbers WITHOUT dollar signs or commas. If not found, use null.",
                    "Extract the total debits, total credits, and proposed new home price from these PSE pages.",
                    pages_to_check,
                )
                pdf_totals = parse_json_from_llm(raw)
                analysis["pdf_totals"] = pdf_totals

                pdf_price = pdf_totals.get("proposed_new_home_price")
                if pdf_price is not None:
                    analysis["pdf_proposed_price"] = pdf_price
                    excel_price = excel_totals["proposed_price"]
                    price_diff = abs(excel_price - pdf_price)
                    analysis["price_difference"] = price_diff

                    if price_diff > 100:
                        issues.append(
                            f"CRITICAL: PSE price mismatch — Signed PDF total is ${pdf_price:,.0f} "
                            f"but PSE Excel total is ${excel_price:,.0f} "
                            f"(difference: ${price_diff:,.0f}). "
                            f"The signed document and spreadsheet must match."
                        )
                    else:
                        analysis["prices_match"] = True
        except Exception as e:
            warnings.append(f"Could not compare PSE PDF vs Excel totals: {str(e)}")

    # === CRITICAL CHECK: Price sheet period vs deposit date ===
    excel_period = excel_totals.get("pricing_period") or ""
    deposit_date_str = (results or {}).get("deposit_date", "") or ""
    print(f"[DEBUG] check_pse_excel: excel_period='{excel_period}', deposit_date_str='{deposit_date_str}'")

    if excel_period and deposit_date_str:
        analysis["excel_pricing_period"] = excel_period
        analysis["deposit_date"] = deposit_date_str

        # Parse the deposit date to get the month
        MONTH_MAP = {
            "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
            "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
        }
        deposit_month = None
        # Try DD/MM/YY or DD/MM/YYYY format
        date_match = re.search(r'(\d{1,2})/(\d{1,2})/(\d{2,4})', deposit_date_str)
        if date_match:
            deposit_month = int(date_match.group(2))

        # Parse the pricing period months (e.g., "Dec-Jan")
        period_months = []
        for m_name, m_num in MONTH_MAP.items():
            if m_name in excel_period.lower():
                period_months.append(m_num)

        if deposit_month and period_months:
            # Check if deposit month falls within the pricing period
            # Handle wrap-around (e.g., Dec-Jan = months 12, 1)
            deposit_in_period = deposit_month in period_months
            analysis["deposit_month"] = deposit_month
            analysis["period_months"] = period_months

            if not deposit_in_period:
                period_names = excel_period
                month_name = list(MONTH_MAP.keys())[deposit_month - 1].title()
                issues.append(
                    f"CRITICAL: Wrong price sheet — PSE Excel is for {period_names} pricing "
                    f"but the deposit was paid in {month_name} (date: {deposit_date_str}). "
                    f"The PSE must use the price sheet that was current when the deposit was paid."
                )
            else:
                analysis["price_sheet_period_valid"] = True

    return {"issues": issues, "warnings": warnings, "analysis": analysis}


# ---------------------------------------------------------------------------
# Detect deposit type from content (not just filenames)
# ---------------------------------------------------------------------------
def detect_deposit_type(file_map: dict, files: list[dict]) -> str:
    """Detect NHP vs STC from file content and names."""
    # Check ITP content if available
    if "itp" in file_map:
        itp_text = extract_pdf_text(file_map["itp"]["full_path"], max_pages=3)
        itp_lower = itp_text.lower()
        if "nhp" in itp_lower or "$2,500" in itp_lower or "$2500" in itp_lower:
            return "NHP"
        if "stc" in itp_lower or "$4,000" in itp_lower or "$4000" in itp_lower:
            return "STC"

    # Check deposit receipt content
    if "deposit_receipt" in file_map:
        receipt_text = extract_pdf_text(file_map["deposit_receipt"]["full_path"], max_pages=2)
        receipt_lower = receipt_text.lower()
        if "2,500" in receipt_lower or "2500" in receipt_lower:
            return "NHP"
        if "4,000" in receipt_lower or "4000" in receipt_lower:
            return "STC"

    # Fallback: check all filenames
    for f in files:
        fn_lower = f["name"].lower()
        if "nhp" in fn_lower or "2500" in fn_lower or "2,500" in fn_lower:
            return "NHP"
        if "stc" in fn_lower or "4000" in fn_lower or "4,000" in fn_lower:
            return "STC"

    return "UNKNOWN"


# ---------------------------------------------------------------------------
# Build corrected zip
# ---------------------------------------------------------------------------
def build_corrected_zip(extract_dir: str, deal_code: str, output_dir: str) -> str:
    zip_name = f"{deal_code}_corrected.zip" if deal_code else "corrected.zip"
    zip_path = os.path.join(output_dir, zip_name)

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fn in sorted(os.listdir(extract_dir)):
            fp = os.path.join(extract_dir, fn)
            if os.path.isfile(fp) and not fn.startswith(".") and not fn.endswith(".md"):
                zf.write(fp, fn)

    return zip_path


# ---------------------------------------------------------------------------
# Helper: get false positive notes for a check
# ---------------------------------------------------------------------------
def _get_fp_notes(check_name: str) -> str:
    try:
        false_positives = db.get_false_positives()
        fps = [fp for fp in false_positives if fp["check_name"] == check_name]
        if fps:
            notes = "\n\nPREVIOUS FALSE POSITIVES TO AVOID (staff-confirmed):\n"
            for fp in fps[:8]:
                notes += f"- Issue '{fp['issue_text']}' was marked incorrect. Note: {fp['notes']}\n"
            return notes
    except Exception:
        pass
    return ""


# ---------------------------------------------------------------------------
# Generate verdict
# ---------------------------------------------------------------------------
def generate_verdict(all_results: dict) -> dict:
    results_summary = json.dumps(all_results, indent=2, default=str)
    consultant_name = all_results.get("consultant_name", "Consultant")
    deal_code = all_results.get("deal_code", "")

    system_prompt = """You are the AUSMAR PSE QA Review Agent. Based on the check results, generate the final QA output.

DECISION FRAMEWORK — TWO STATES ONLY:
- ACCEPTED: Submission passes QA. May include warnings/notes but no blocking issues.
- NOT ACCEPTED: One or more blocking issues found. Consultant must resubmit.

BLOCKING ISSUES (verdict = NOT ACCEPTED):
- Missing a CORE REQUIRED document (PSE Doc, PSE Excel, GeoSite, ITP, Deposit Receipt, Drivers Licence, Red Pen, Promo Ack)
- PSE Excel total and PSE Signed PDF total differ by more than $100
- PSE price sheet period doesn't match the deposit date month
- Any issue containing the word "CRITICAL" in the check results
- Plan demonstrably does not fit the lot (width exceeds available space after setbacks)

NON-BLOCKING (include in warnings, do NOT reject for these alone):
- Setback minimums close to limit but within tolerance
- Unclassified files (unless they are clearly a missing required doc)
- Admin/naming issues (already auto-fixed)
- LHDC checks (disabled — not a blocking issue)

CRITICAL RULES:
1. Gas cooktops are EXEMPT from covenant gas bans per Stockland ruling. Do NOT flag them.
2. Red Pen markups using colored highlights (green, yellow, blue, pink) ARE valid — do NOT flag them as "not in red".
3. ACCURACY IS CRITICAL: Do NOT flag false positives. Only flag issues you are confident about.
4. Admin/naming corrections have been auto-applied. Don't mention those as issues.

Output valid JSON:
{
  "verdict": "ACCEPTED or NOT ACCEPTED",
  "verdict_reason": "one-line reason",
  "critical_issues": ["list of blocking issues only — empty if ACCEPTED"],
  "warnings": ["list of non-blocking concerns"],
  "heath_review_note": "Technical summary for Heath Nunn (Drafting Manager). Cover GeoSite accuracy, Red Pen quality, plan-to-lot fit, site coverage, acoustic/covenant concerns. Include PSE price comparison results if available. Be specific with numbers. 2-4 paragraphs.",
  "consultant_feedback_email": "Professional email to the consultant. Start with 'Hi [consultant name],' and end with 'Regards,\\nNik'. Be specific about what needs fixing. If accepted, acknowledge good work briefly."
}"""

    raw = call_text_model(system_prompt, f"QA results for {deal_code}:\n\n{results_summary}", model="gpt-4.1-mini")
    result = parse_json_from_llm(raw)

    # Normalise verdict to exactly two states — collapse any LLM drift
    v = str(result.get("verdict", "")).upper()
    if "NOT ACCEPTED" in v or "NOT ACCEPT" in v:
        result["verdict"] = "NOT ACCEPTED"
    else:
        result["verdict"] = "ACCEPTED"

    return result


# ---------------------------------------------------------------------------
# Cross-check against pre-logged job info
# ---------------------------------------------------------------------------
def cross_check_prelog(prelog: dict, review_results: dict) -> list[str]:
    notes = []

    if prelog.get("deposit_amount"):
        amt = prelog["deposit_amount"]
        dep_type = review_results.get("deposit_type", "")
        if amt == 2500 and dep_type == "STC":
            notes.append(
                "Pre-log says $2,500 (NHP) but submission appears to be STC — "
                "known pattern: wholesale groups sometimes take larger deposits"
            )
        elif amt == 4000 and dep_type == "NHP":
            notes.append(
                "Pre-log says $4,000 (STC) but submission appears to be NHP — "
                "verify correct ITP report used"
            )

    if prelog.get("consultant_name"):
        gs_analysis = review_results.get("checks", {}).get("geosite_verification", {}).get("analysis", {})
        gs_consultant = gs_analysis.get("consultant_name", "")
        if gs_consultant and prelog["consultant_name"].lower() not in gs_consultant.lower():
            notes.append(
                f"Pre-log consultant '{prelog['consultant_name']}' doesn't match "
                f"GeoSite consultant '{gs_consultant}'"
            )

    return notes


# ---------------------------------------------------------------------------
# Main QA Review Pipeline (v2: Content-Based)
# ---------------------------------------------------------------------------
def run_qa_review(zip_path: str, zip_name: str, corrected_zip_dir: str,
                  progress_callback=None, notes: str = "") -> dict:
    """Run the full QA review with content-based document classification.
    
    Args:
        notes: Pre-Log information (e.g., Land Registered Y/N, NHP/STC, etc.)
    """

    def _progress(pct, msg):
        if progress_callback:
            try:
                progress_callback(pct, msg)
            except Exception:
                pass

    # Parse Pre-Log notes
    prelog = parse_prelog(notes)

    results = {
        "zip_name": zip_name,
        "timestamp": datetime.now().isoformat(),
        "deposit_type": "UNKNOWN",
        "deal_code": "",
        "checks": {},
        "corrections_applied": [],
        "consultant_name": "",
        "prelog": prelog,
    }

    extract_dir = tempfile.mkdtemp(prefix="ausmar_qa_")

    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            zf.extractall(extract_dir)
    except zipfile.BadZipFile:
        return {"error": "Invalid zip file — could not extract"}

    # Remove __MACOSX
    macosx = os.path.join(extract_dir, "__MACOSX")
    if os.path.exists(macosx):
        shutil.rmtree(macosx)

    try:
        # === Check 1: File Structure (flatten, remove junk) ===
        _progress(5, "Checking file structure...")
        structure = check_file_structure(extract_dir, zip_name)
        results["checks"]["file_structure"] = {
            "issues": structure["issues"],
            "warnings": structure["warnings"],
            "file_count": len(structure["files"]),
            "files": [f["name"] for f in structure["files"]],
        }
        results["corrections_applied"] = structure["corrections"]
        files = structure["files"]
        deal_code = structure["zip_stem"]
        results["deal_code"] = deal_code

        # === NEW: Content-Based Document Classification ===
        _progress(10, "Classifying documents by content (text extraction)...")
        classification_result = classify_all_files(files)
        file_map = classification_result["file_map"]
        unclassified = classification_result["unclassified"]
        classifications = classification_result["classifications"]

        _progress(25, "Document classification complete. Checking completeness...")

        # Store classification info in results
        results["checks"]["document_classification"] = {
            "classifications": {
                fname: {
                    "identified_as": DOC_TYPES.get(info.get("doc_type", ""), "Unknown"),
                    "confidence": info.get("confidence", 0),
                    "method": info.get("method", ""),
                }
                for fname, info in classifications.items()
                if info.get("doc_type")
            },
            "unclassified_files": [f["name"] for f in unclassified],
        }

        # === Rename files to standard names ===
        rename_corrections = rename_classified_files(extract_dir, file_map, classifications)
        results["corrections_applied"].extend(rename_corrections)

        # === Detect deposit type from content ===
        results["deposit_type"] = detect_deposit_type(file_map, files)

        # === Pre-Log File Merge ===
        # Look up Pre-Log by deal code. If files were uploaded there (ITP, receipt, licence),
        # copy any that are missing from the zip into extract_dir so they appear in the
        # corrected zip and don't block the submission as "missing".
        _progress(28, "Checking Pre-Log for supplementary files...")
        try:
            prelog_db = db.find_prelog_by_deal_code(deal_code)
        except Exception:
            prelog_db = None
        prelog_merged_docs = []
        if prelog_db and prelog_db.get("file_paths"):
            for fp in prelog_db["file_paths"]:
                if not os.path.isfile(fp):
                    continue
                fn = os.path.basename(fp)
                dest = os.path.join(extract_dir, fn)
                # Only copy if a file with that name isn't already there
                if not os.path.exists(dest):
                    shutil.copy2(fp, dest)
                    prelog_merged_docs.append(fn)
                    results["corrections_applied"].append(f"Merged from Pre-Log: {fn}")
            if prelog_merged_docs:
                # Re-classify now that we have extra files
                new_files = []
                for fn in prelog_merged_docs:
                    fp2 = os.path.join(extract_dir, fn)
                    ext = os.path.splitext(fn)[1].lower()
                    new_files.append({"name": fn, "full_path": fp2, "ext": ext, "size": os.path.getsize(fp2)})
                extra_result = classify_all_files(new_files)
                # Merge into existing file_map (don't overwrite already-classified docs)
                for doc_type, finfo in extra_result["file_map"].items():
                    if doc_type not in file_map:
                        file_map[doc_type] = finfo
                        classifications[finfo["name"]] = extra_result["classifications"].get(finfo["name"], {})
                results["prelog_merged_files"] = prelog_merged_docs
                results["prelog_id"] = prelog_db["id"]

        # === Check 2: Document Completeness (content-based) ===
        _progress(30, "Checking document completeness...")
        completeness = check_document_completeness(file_map, unclassified)
        results["checks"]["document_completeness"] = {
            "issues": completeness["issues"],
            "warnings": completeness["warnings"],
            "found_documents": list(completeness["found"].keys()),
            "found_details": completeness["found"],  # maps canonical name -> original filename
            "conditional_documents": completeness["conditional_found"],
        }

        # === Check 3: GeoSite Verification (Vision) ===
        _progress(40, "Analysing GeoSite with vision AI...")
        # Determine land registration status from Pre-Log notes or DB prelog notes
        _land_registered = prelog.get("land_registered")  # from parse_prelog(notes)
        if _land_registered is None and prelog_db:
            _db_notes = (prelog_db.get("notes") or "").lower()
            if "land registered: y" in _db_notes or "land registered: yes" in _db_notes:
                _land_registered = True
            elif "land registered: n" in _db_notes or "land registered: no" in _db_notes:
                _land_registered = False
        geosite_result = check_geosite(file_map, land_registered=_land_registered)
        results["checks"]["geosite_verification"] = {
            "issues": geosite_result["issues"],
            "warnings": geosite_result["warnings"],
            "analysis": geosite_result.get("analysis", {}),
            "lot_dimensions": geosite_result.get("lot_dimensions"),
        }

        gs_analysis = geosite_result.get("analysis", {})

        # === Extract consultant name and key data from ITP FIRST (most reliable source) ===
        # The ITP has a grey header bar: "PURCHASER/S DETAILS" on left, "CONSULTANT" label
        # in the middle, and the consultant name (e.g., "TELFORD LOUEZ") on the right.
        # The purchaser names are BELOW in the form fields — they are NOT the consultant.
        if "itp" in file_map:
            try:
                itp_path = file_map["itp"]["full_path"]
                itp_pages = pdf_all_pages_to_base64(itp_path, dpi=150, max_pages=3)
                if itp_pages:
                    raw = call_vision_model(
                        "You are extracting data from an AUSMAR Intention to Purchase (ITP) form. "
                        "CRITICAL LAYOUT INFORMATION:\n"
                        "- There is a GREY HEADER BAR near the top of page 1.\n"
                        "- On the LEFT side of this grey bar: 'PURCHASER/S DETAILS' label.\n"
                        "- On the RIGHT side of this grey bar: 'CONSULTANT' label followed by the consultant's name in BOLD.\n"
                        "- The CONSULTANT name is the AUSMAR sales person (e.g., 'TELFORD LOUEZ').\n"
                        "- BELOW the grey bar are the purchaser form fields (First Name, Middle Name, Surname).\n"
                        "- The purchaser names are the BUYERS, NOT the consultant.\n"
                        "- DO NOT confuse purchaser names with the consultant name.\n\n"
                        "Also extract from page 1: lot number, street name, estate/suburb, land price.\n"
                        "From the deposit section (page 1 or 2): the deposit amount.\n"
                        "From the PURCHASER SIGNATURE section (usually page 3): the DATE next to the purchaser signatures.\n"
                        "This is the deposit/signing date in DD/MM/YYYY or DD/MM/YY format (e.g., 15/02/2026).\n\n"
                        "Return ONLY a JSON object:\n"
                        '{"consultant_name": "Full Name from grey header bar", '
                        '"purchaser_names": ["Purchaser 1 Full Name", "Purchaser 2 Full Name"], '
                        '"lot_number": "number", "street": "street name", "estate": "estate or suburb name", '
                        '"deposit_date": "DD/MM/YYYY as written on the form", '
                        '"land_price": "dollar amount", '
                        '"deposit_amount": "dollar amount"}. '
                        "If a field is not visible, use null.",
                        "Look at the GREY HEADER BAR at the top. The CONSULTANT name is on the RIGHT side of that bar, in bold. "
                        "The purchaser names are in the form fields BELOW the bar. Extract all requested data.",
                        itp_pages,
                    )
                    itp_info = parse_json_from_llm(raw)
                    if itp_info.get("consultant_name"):
                        results["consultant_name"] = itp_info["consultant_name"]
                    if itp_info.get("purchaser_names"):
                        results["purchaser_names"] = itp_info["purchaser_names"]
                    if itp_info.get("deposit_date"):
                        results["deposit_date"] = itp_info["deposit_date"]
                    if itp_info.get("lot_number"):
                        results["lot_number"] = itp_info["lot_number"]
                    if itp_info.get("street"):
                        results["street"] = itp_info["street"]
                    if itp_info.get("estate"):
                        results["estate"] = itp_info["estate"]
                    if itp_info.get("land_price"):
                        results["land_price"] = itp_info["land_price"]
                    if itp_info.get("deposit_amount"):
                        results["deposit_amount"] = itp_info["deposit_amount"]
            except Exception as e:
                print(f"[WARN] ITP data extraction failed: {e}")

        # If deposit_date still not set after ITP extraction, try to extract from ITP page 3 specifically
        if not results.get("deposit_date") and "itp" in file_map:
            try:
                itp_path = file_map["itp"]["full_path"]
                # Get just page 3 at higher DPI for the signature date
                page3 = pdf_page_to_base64(itp_path, page_num=2, dpi=150)  # 0-indexed, so page 3 = index 2
                if page3:
                    raw = call_vision_model(
                        "This is the signature page of an AUSMAR ITP form. "
                        "Find the DATE written next to the PURCHASER SIGNATURE. "
                        "It will be in DD/MM/YYYY or DD/MM/YY format (e.g., 15/02/2026). "
                        "Return ONLY a JSON object: {\"deposit_date\": \"DD/MM/YYYY\"}",
                        "Extract the date from the purchaser signature section.",
                        [page3],
                    )
                    date_info = parse_json_from_llm(raw)
                    if date_info.get("deposit_date"):
                        results["deposit_date"] = date_info["deposit_date"]
                        print(f"[DEBUG] deposit_date extracted from ITP page 3: {results['deposit_date']}")
            except Exception as e:
                print(f"[WARN] deposit_date fallback extraction failed: {e}")

        # Fallback: GeoSite consultant name (only if ITP didn't provide one)
        gs_consultant = gs_analysis.get("consultant_name", "")
        if not results["consultant_name"] and gs_consultant:
            results["consultant_name"] = gs_consultant

        # Fallback 2: PSE Doc page 1 — has "Salesperson: Name" field
        if not results["consultant_name"] and "pse_doc" in file_map:
            try:
                pse_path = file_map["pse_doc"]["full_path"]
                pse_pages = pdf_all_pages_to_base64(pse_path, dpi=100, max_pages=1)
                if pse_pages:
                    raw = call_vision_model(
                        "You are extracting data from an AUSMAR Provisional Sales Estimate (PSE) document. "
                        "The Salesperson name is near the top of page 1, labelled 'Salesperson:'. "
                        "This is the consultant/sales person who created the PSE. "
                        "Do NOT confuse with the client/purchaser name. "
                        "Return ONLY a JSON object: "
                        '{"salesperson": "Full Name", "client_name": "Full Name", '
                        '"job_code": "code", "site_address": "address", "pricing_period": "month-month year"}. '
                        "If a field is not visible, use null.",
                        "Extract the salesperson name, client name, job code, site address, and pricing period.",
                        pse_pages,
                    )
                    pse_info = parse_json_from_llm(raw)
                    if pse_info.get("salesperson"):
                        results["consultant_name"] = pse_info["salesperson"]
                    if pse_info.get("client_name"):
                        results["client_name"] = pse_info["client_name"]
                    if pse_info.get("job_code") and not results.get("deal_code"):
                        results["deal_code"] = pse_info["job_code"]
                    if pse_info.get("site_address"):
                        results["site_address"] = pse_info["site_address"]
                    if pse_info.get("pricing_period"):
                        results["pricing_period"] = pse_info["pricing_period"]
            except Exception as e:
                print(f"[WARN] Consultant name extraction from PSE failed: {e}")

        # === Check 4: Plan-to-Lot Fit ===
        _progress(55, "Verifying plan-to-lot fit...")
        fit_result = check_plan_to_lot_fit(geosite_result, file_map)
        results["checks"]["plan_to_lot_fit"] = {
            "issues": fit_result["issues"],
            "warnings": fit_result["warnings"],
            "plan_identified": fit_result.get("plan_identified"),
            "fit_result": fit_result.get("fit_result"),
            "lot_dimensions": fit_result.get("lot_dimensions"),
            "plan_specs": fit_result.get("plan_specs"),
        }

        # === Check 5: Red Pen Markup (Vision) ===
        _progress(70, "Analysing Red Pen markups...")
        dep_type = results["deposit_type"] if results["deposit_type"] != "UNKNOWN" else "NHP"
        redpen_result = check_red_pen(file_map, dep_type)
        results["checks"]["red_pen_markup"] = {
            "issues": redpen_result["issues"],
            "warnings": redpen_result["warnings"],
            "analysis": redpen_result.get("analysis", {}),
        }

        # === Check 6: PSE Excel ===
        _progress(82, "Checking PSE patterns...")
        pse_result = check_pse_excel(file_map, gs_analysis, results)
        results["checks"]["pse_analysis"] = {
            "issues": pse_result["issues"],
            "warnings": pse_result["warnings"],
            "analysis": pse_result.get("analysis", {}),
        }

        # === Check 7: Sites with Fall (conditional) ===
        fall_mm = gs_analysis.get("fall_across_site_mm")
        has_fall_ack = "fall_ack" in file_map
        if fall_mm and isinstance(fall_mm, (int, float)) and fall_mm >= 500 and not has_fall_ack:
            if "geosite_verification" in results["checks"]:
                results["checks"]["geosite_verification"]["warnings"].append(
                    f"Site fall {fall_mm}mm >= 500mm but no Sites with Fall Acknowledgment (Signed) found in submission"
                )

        # === Check 8: Red Pen to PSE Cross-Reference ===
        _progress(85, "Cross-referencing Red Pen annotations with PSE...")
        redpen_analysis = results["checks"].get("red_pen_markup", {}).get("analysis", {})
        ref_numbers = redpen_analysis.get("reference_numbers_found", [])
        if ref_numbers and "pse_doc" in file_map:
            try:
                pse_path = file_map["pse_doc"]["full_path"]
                # Check PSE pages 4-10 for reference number write-ups (sections are in the middle)
                # Use page-by-page to save memory
                pse_pages_xref = []
                for pg in range(3, 10):  # pages 4-10 (0-indexed: 3-9)
                    b64 = pdf_page_to_base64(pse_path, page_num=pg, dpi=100)
                    if b64:
                        pse_pages_xref.append(b64)
                if pse_pages_xref:
                    refs_str = ", ".join(str(r) for r in ref_numbers)
                    raw = call_vision_model(
                        "You are checking an AUSMAR Provisional Sales Estimate (PSE) document. "
                        "The Red Pen markup has the following reference numbers annotated: "
                        f"{refs_str}. "
                        "Check if the PSE document contains write-ups or sections that correspond "
                        "to these reference numbers. PSE sections are typically numbered like "
                        "3.0, 3.1, 3.2, 3.2.a etc. and describe structural changes, additions, or modifications.\n\n"
                        "Return ONLY a JSON object:\n"
                        '{"references_found_in_pse": ["list of ref numbers that have matching PSE sections"], '
                        '"references_missing_from_pse": ["list of ref numbers NOT found in PSE"], '
                        '"pse_sections_found": ["list of PSE section numbers found in the document"], '
                        '"notes": "brief description"}',
                        "Check if these Red Pen reference numbers have corresponding write-ups in the PSE document.",
                        pse_pages_xref[:5],  # Max 5 pages to keep within token limits
                    )
                    xref_result = parse_json_from_llm(raw)
                    results["checks"]["redpen_pse_crossref"] = {
                        "issues": [],
                        "warnings": [],
                        "analysis": xref_result,
                    }
                    missing_refs = xref_result.get("references_missing_from_pse", [])
                    if missing_refs:
                        results["checks"]["redpen_pse_crossref"]["warnings"].append(
                            f"Red Pen reference numbers {', '.join(str(r) for r in missing_refs)} "
                            f"may not have corresponding write-ups in the PSE document. "
                            f"Verify all Red Pen annotations are documented in the PSE."
                        )
            except Exception as e:
                print(f"[WARN] Red Pen to PSE cross-reference failed: {e}")

        # === Check 9: LHDC — DISABLED (100% false positive rate on 14/14 staff reviews) ===
        # LHDC checks require physical inspection of stamped drawings, not floor plan PDFs.
        # Re-enable only if AUSMAR confirms a specific council mandates it at PSE stage.
        _progress(88, "LHDC check skipped (disabled)...")
        # Still detect regional rules for other checks (setbacks, coverage)
        estate_name = results.get("estate") or gs_analysis.get("estate_name") or ""
        street_addr = results.get("street") or gs_analysis.get("street_address") or ""
        regional_rules = detect_regional_rules(estate_name, street_addr)
        results["regional_rules"] = regional_rules["name"] if regional_rules else "Unknown (no council match)"
        results["lhdc_applicable"] = False
        results["lhdc_reason"] = "LHDC check disabled — not applicable at PSE submission stage"
        results["checks"]["lhdc_assessment"] = {
            "issues": [],
            "warnings": [],
            "analysis": {"skipped": True, "reason": "LHDC check disabled — not applicable at PSE submission stage"},
            "lhdc_reason": "LHDC check disabled — not applicable at PSE submission stage",
        }

        # === Build corrected zip (always build so user can download renamed files) ===
        corrected_path = build_corrected_zip(extract_dir, deal_code, corrected_zip_dir)
        results["corrected_zip_path"] = corrected_path
        results["corrected_zip_filename"] = os.path.basename(corrected_path)

        # === Cross-check pre-log ===
        _progress(92, "Cross-checking pre-log data...")
        # Use already-fetched prelog_db from merge step (avoid double DB lookup)
        try:
            prelog = prelog_db if "prelog_db" in dir() else db.find_prelog_by_deal_code(deal_code)
        except Exception:
            prelog = None
        prelog_notes = []
        if prelog:
            prelog_notes = cross_check_prelog(prelog, results)
            results["prelog_id"] = prelog["id"]
            results["prelog_notes"] = prelog_notes
            if prelog_notes:
                results["checks"]["prelog_crosscheck"] = {
                    "issues": [], "warnings": prelog_notes,
                }

        # === Generate verdict ===
        _progress(95, "Generating verdict and outputs...")
        verdict = generate_verdict(results)
        results["verdict_data"] = verdict

        _progress(98, "Finalising...")

    except Exception as e:
        results["error"] = f"Review pipeline error: {str(e)}\n{traceback.format_exc()}"
    finally:
        shutil.rmtree(extract_dir, ignore_errors=True)
        gc.collect()

    return results
